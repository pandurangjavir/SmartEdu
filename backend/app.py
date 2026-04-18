from flask import Flask, request, jsonify
from flask_cors import CORS
from werkzeug.utils import secure_filename
from models import (db, User, Department, Class, Subject, Student, Fee, Mark, 
                   Attendance, Event, EventRegistration, Announcement, Course, Notification, ChatMessage,
                   Fine, get_fee_details, get_attendance, get_upcoming_events)
from rasa_service import rasa_service
from config import config
import json
import os
from datetime import datetime, timedelta, date
from decimal import Decimal
from textblob import TextBlob
import pypdf
import io
import docx
import re
import jwt

app = Flask(__name__)
CORS(app, origins=['http://localhost:3000', 'http://localhost:3001', 'http://127.0.0.1:3000', 'http://127.0.0.1:3001'])

# Load configuration
config_name = os.getenv('FLASK_ENV', 'development')
app.config.from_object(config[config_name])

# Initialize database
db.init_app(app)


def get_current_user_from_token():
    """
    Helper to extract current user from bearer token.
    Returns (user, None) on success or (None, (response, status_code)) on error.
    """
    auth_header = request.headers.get('Authorization')
    if not auth_header or not auth_header.startswith('Bearer '):
        return None, (jsonify({'error': 'Authorization token required'}), 401)

    token = auth_header.split(' ')[1]
    
    try:
        secret_key = app.config.get('SECRET_KEY', 'dev_secret_key_12345')
        payload = jwt.decode(token, secret_key, algorithms=['HS256'])
        user_id = payload.get('user_id')
    except jwt.ExpiredSignatureError:
        return None, (jsonify({'error': 'Token has expired'}), 401)
    except jwt.InvalidTokenError:
        return None, (jsonify({'error': 'Invalid authentication token'}), 401)

    user = db.session.get(User, user_id)
    if not user:
        return None, (jsonify({'error': 'User not found'}), 404)

    return user, None

def detect_admin_filters(message):
    """Detect class, roll number, or name filters from message"""
    filters = {
        'class_filter': None,
        'roll_filter': None,
        'name_filter': None,
        'target_student_id': None
    }
    
    # First check for specific roll number match
    roll_search = Student.query.filter_by(roll_no=message).first()
    if roll_search:
        filters['target_student_id'] = roll_search.student_id
        filters['roll_filter'] = message
        return filters
    
    # Check for partial roll number match
    roll_search = Student.query.filter(Student.roll_no.ilike(f'%{message}%')).first()
    if roll_search:
        filters['target_student_id'] = roll_search.student_id
        filters['roll_filter'] = message
        return filters
    
    words = message.lower().split()
    for word in words:
        # Check for class filters (TY, SY, FY, BE with optional branch like CSE, ECE, etc.)
        word_upper = word.upper()
        # Match patterns like TY, SY, FY, BE, or combination like TY-CSE, SY-CSE, etc.
        class_patterns = ['TY-CSE', 'TY-ECE', 'TY-IT', 'SY-CSE', 'SY-ECE', 'SY-IT', 'FY-CSE', 'FY-ECE', 'FY-IT']
        if any(pattern in word_upper for pattern in class_patterns) or word_upper in ['TY', 'SY', 'FY', 'BE']:
            filters['class_filter'] = word
        # Check for roll number (numbers with 2+ digits)
        elif word.isdigit() and len(word) >= 2:
            filters['roll_filter'] = word
        # Check for student name
        else:
            student_search = Student.query.join(User).filter(User.name.ilike(f'%{word}%')).first()
            if student_search:
                filters['name_filter'] = word
                filters['target_student_id'] = student_search.student_id
                break
    
    # If roll number is specified, find student by roll number
    if filters['roll_filter'] and not filters['target_student_id']:
        student_search = Student.query.filter_by(roll_no=filters['roll_filter']).first()
        if student_search:
            filters['target_student_id'] = student_search.student_id
    
    return filters

# Authentication endpoints
@app.route('/api/auth/login', methods=['POST'])
def login():
    """
    Authenticate user login
    """
    try:
        data = request.get_json()
        
        if not data or not data.get('email') or not data.get('password'):
            return jsonify({'error': 'Email and password are required'}), 400
        
        email = data['email']
        password = data['password']
        
        # Find user by email
        user = User.query.filter_by(email=email).first()
        
        if not user:
            return jsonify({'error': 'Invalid email or password'}), 401
        
        # Check password securely
        if not user.check_password(password):
            return jsonify({'error': 'Invalid email or password'}), 401
        
        # Generate secure JWT token
        secret_key = app.config.get('SECRET_KEY', 'dev_secret_key_12345')
        payload = {
            'user_id': user.user_id,
            'role': user.role,
            'exp': datetime.utcnow() + timedelta(days=7) # Token valid for 7 days
        }
        token = jwt.encode(payload, secret_key, algorithm='HS256')
        
        # Return user data and token
        user_data = {
            'user_id': user.user_id,
            'name': user.name,
            'email': user.email,
            'role': user.role,
            'contact_no': user.contact_no
        }
        
        return jsonify({
            'success': True,
            'token': token,
            'user': user_data
        }), 200
        
    except Exception as e:
        return jsonify({'error': f'Login failed: {str(e)}'}), 500

@app.route('/api/auth/profile', methods=['GET'])
def get_profile():
    """
    Get user profile (requires authentication)
    """
    try:
        user, error = get_current_user_from_token()
        if error:
            return error

        user_data = {
            'user_id': user.user_id,
            'name': user.name,
            'email': user.email,
            'role': user.role,
            'contact_no': user.contact_no
        }
        
        return jsonify({'user': user_data}), 200
            
    except Exception as e:
        return jsonify({'error': f'Profile fetch failed: {str(e)}'}), 500

# Sentiment Analysis Functions
def analyze_sentiment(text):
    """
    Analyze sentiment of text using TextBlob
    Returns polarity score (-1 to 1) and sentiment label
    """
    try:
        blob = TextBlob(text)
        polarity = blob.sentiment.polarity
        
        # Determine sentiment label
        if polarity > 0.1:
            sentiment_label = "positive"
        elif polarity < -0.1:
            sentiment_label = "negative"
        else:
            sentiment_label = "neutral"
            
        return {
            'polarity': polarity,
            'sentiment': sentiment_label
        }
    except Exception as e:
        print(f"Sentiment analysis error: {str(e)}")
        return {
            'polarity': 0.0,
            'sentiment': 'neutral'
        }

def get_empathetic_prefix(sentiment_data):
    """
    Get empathetic message prefix based on sentiment
    """
    if sentiment_data['polarity'] < 0:
        empathetic_messages = [
            "I understand this might be frustrating. ",
            "I can see you're having some concerns. ",
            "I'm here to help with whatever is troubling you. ",
            "I sense you might be feeling stressed about this. ",
            "I want to make sure I address your concerns properly. "
        ]
        import random
        return random.choice(empathetic_messages)
    return ""

def get_educational_response(intent, parameters, user_message, fulfillment_text=None):
    """
    Generate educational responses based on detected intent.
    Mirrors previous behavior without Dialogflow dependency.
    """
    # Prefer fulfillment_text if provided by NLU
    if fulfillment_text and str(fulfillment_text).strip():
        base_response = fulfillment_text
    else:
        intent_responses = {
            'course_inquiry': 'I can help you with course information! We offer various programming courses including Python, JavaScript, and Web Development. Would you like to know more about any specific course?',
            'event_inquiry': 'Great! We have several upcoming events. Let me show you what\'s available.',
            'help_request': 'I\'m here to help! I can assist you with course information, event details, announcements, and general questions about our educational platform.',
            'greeting': 'Hello! Welcome to SmartEdu! I\'m your educational assistant. How can I help you today?',
            'goodbye': 'Thank you for using SmartEdu! Have a great day and feel free to come back anytime for assistance.',
            'Default Fallback Intent': 'I understand you\'re looking for help. I can assist you with course information, upcoming events, announcements, or answer general questions about our educational platform.'
        }
        base_response = intent_responses.get(intent, intent_responses['Default Fallback Intent'])

    intent_suggestions = {
        'course_inquiry': ['Python Programming', 'Web Development', 'Database Design', 'View All Courses'],
        'event_inquiry': ['View Events', 'Register for Event', 'Event Details', 'Upcoming Workshops'],
        'help_request': ['Course Information', 'Events', 'Announcements', 'Contact Support'],
        'greeting': ['Course Information', 'Events', 'Help', 'About SmartEdu'],
        'goodbye': [],
        'Default Fallback Intent': ['Course Information', 'Events', 'Announcements', 'Help']
    }
    suggestions = intent_suggestions.get(intent, intent_suggestions['Default Fallback Intent'])
    return {
        'response': base_response,
        'suggestions': suggestions
    }

# --------- College Information Response Function ---------

# Removed redundant get_college_info_response as we now use Gemini with Admission.json context


# Removed redundant logic as we now use Gemini for college info and guidance


# --------- Rasa Chatbot minimal endpoint (MySQL-backed) ---------
@app.route('/chatbot', methods=['GET'])
def chatbot_health():
    # Allow simple GET probes to avoid 405 noise in console
    return jsonify({"ok": True, "endpoint": "/chatbot", "method": "POST"}), 200

@app.route('/api/chatbot/voice', methods=['POST'])
def chatbot_voice():
    """Handle voice messages from the chatbot with full speech-to-text processing"""
    try:
        # Check if audio file is present
        if 'audio' not in request.files:
            return jsonify({'error': 'No audio file provided'}), 400
        
        audio_file = request.files['audio']
        if audio_file.filename == '':
            return jsonify({'error': 'No audio file selected'}), 400
        
        # Get user_id from request
        user_id = request.form.get('user_id')
        student_id = request.form.get('student_id')
        
        # Save the audio file temporarily
        import tempfile
        import os
        with tempfile.NamedTemporaryFile(delete=False, suffix='.wav') as temp_file:
            audio_file.save(temp_file.name)
            temp_audio_path = temp_file.name
        
        try:
            import requests
            
            # Deepgram STT
            deepgram_api_key = os.environ.get('DEEPGRAM_API_KEY')
            if not deepgram_api_key:
                return jsonify({'error': 'Deepgram API key not configured'}), 500
                
            stt_url = "https://api.deepgram.com/v1/listen?model=nova-2&smart_format=true"
            headers = {
                "Authorization": f"Token {deepgram_api_key}",
                "Content-Type": "audio/wav"
            }
            
            with open(temp_audio_path, 'rb') as f:
                stt_response = requests.post(stt_url, headers=headers, data=f)
                
            if stt_response.status_code == 200:
                stt_data = stt_response.json()
                try:
                    text = stt_data['results']['channels'][0]['alternatives'][0]['transcript']
                    print(f"Recognized text via Deepgram: {text}")
                except (KeyError, IndexError):
                    text = ""
            else:
                print(f"Deepgram STT error: {stt_response.text}")
                text = ""
                
            if not text.strip():
                response_text = 'Sorry, I could not understand your voice. Please try speaking more clearly.'
                chatbot_data = {'intent': 'voice_error', 'response': response_text, 'data': {}}
            else:
                # Process the recognized text through the chatbot
                chatbot_payload = {
                    'message': text,
                    'user_id': user_id,
                    'student_id': student_id
                }
                
                # Call the main chatbot route
                from flask import g
                with app.test_request_context('/chatbot', method='POST', json=chatbot_payload):
                    chatbot_response = chatbot_route()
                    chatbot_data = chatbot_response.get_json()
                    
            response_text = chatbot_data.get('response', 'Sorry, I could not process your request.')
            
            # Deepgram TTS
            tts_url = "https://api.deepgram.com/v1/speak?model=aura-asteria-en"
            tts_headers = {
                "Authorization": f"Token {deepgram_api_key}",
                "Content-Type": "application/json"
            }
            tts_payload = {"text": response_text}
            
            audio_response_path = temp_audio_path.replace('.wav', '_response.mp3')
            tts_response = requests.post(tts_url, headers=tts_headers, json=tts_payload)
            
            if tts_response.status_code == 200:
                with open(audio_response_path, 'wb') as f:
                    f.write(tts_response.content)
                    
                with open(audio_response_path, 'rb') as audio_file:
                    audio_data = audio_file.read()
                    
                # Clean up temporary files
                os.unlink(temp_audio_path)
                os.unlink(audio_response_path)
                
                return jsonify({
                    'intent': chatbot_data.get('intent', 'voice_message'),
                    'response': response_text,
                    'data': chatbot_data.get('data', {}),
                    'audio_response': True,
                    'audio_data': audio_data.hex()
                })
            else:
                print(f"Deepgram TTS error: {tts_response.text}")
                os.unlink(temp_audio_path)
                return jsonify({
                    'intent': chatbot_data.get('intent', 'voice_message'),
                    'response': response_text,
                    'data': chatbot_data.get('data', {}),
                    'audio_response': False
                })
                
        except Exception as e:
            print(f"Error in voice processing: {e}")
            # Clean up temporary files
            try:
                if os.path.exists(temp_audio_path):
                    os.unlink(temp_audio_path)
                if 'audio_response_path' in locals() and os.path.exists(audio_response_path):
                    os.unlink(audio_response_path)
            except:
                pass
            
            return jsonify({
                'intent': 'voice_error',
                'response': 'Sorry, there was an error processing your voice message. Please try again.',
                'data': {}
            })
        
    except Exception as e:
        print(f"Error processing voice message: {e}")
        return jsonify({'error': 'Failed to process voice message'}), 500

@app.route('/chatbot', methods=['POST'])
def chatbot_route():
    payload = request.get_json(silent=True) or {}
    message = (payload.get('message') or '').strip()
    student_id = payload.get('student_id')
    user_id = payload.get('user_id')
    
    # Determine user role
    user_role = None
    admin_user = None
    if user_id:
        try:
            admin_user = db.session.get(User, int(user_id))
            if admin_user:
                user_role = admin_user.role
        except:
            pass
    
    if not message:
        return jsonify({'error': 'message is required'}), 400

    # Use Hybrid NLP Engine for Intent Detection
    msg = message.lower()
    from nlp_engine import nlp_engine
    intent = 'unknown'
    confidence = 1.0
    # Enhanced keyword matching for all intents - order matters!
    # Check specific intents first before general ones
    
    # Farewell intent (check first to avoid conflicts)
    if any(phrase in msg for phrase in ['bye', 'goodbye', 'see you later', 'catch you later', 'talk to you soon', 'take care', 'see ya', 'good night', 'bye smartedu', 'thanks bye', 'ok bye', 'bye for now']):
        intent = 'goodbye'
    
    # How are you intent
    elif any(phrase in msg for phrase in ['how are you', 'how are you doing', 'how\'s it going', 'how\'s your day', 'how have you been', 'what\'s up smartedu', 'are you fine', 'you good']):
        intent = 'ask_howareyou'
    
    # Who are you intent
    elif any(phrase in msg for phrase in ['who are you', 'what are you', 'tell me about yourself', 'what\'s your name', 'introduce yourself', 'are you smartedu', 'who is smartedu', 'what is smartedu']):
        intent = 'ask_whoareyou'
    
    # Are you bot intent
    elif any(phrase in msg for phrase in ['are you a bot', 'are you a robot', 'are you real', 'are you human', 'are you alive', 'do you have feelings', 'are you ai', 'are you chatbot']):
        intent = 'ask_areyoubot'
    
    # Creator intent
    elif any(phrase in msg for phrase in ['who made you', 'who created you', 'who developed you', 'who built you', 'who designed you', 'tell me your developer name', 'who is your owner']):
        intent = 'ask_creator'
    
    # Help intent
    elif any(phrase in msg for phrase in ['can you help me', 'i need help', 'please help', 'what can you do', 'what services do you provide', 'how can you help me', 'help me smartedu', 'anyone here', 'are you there', 'are you online', 'smartedu can you assist me']):
        intent = 'ask_help'
    
    # Conversational / Identity (Check these carefully to avoid generic greeting fallback)
    elif any(phrase in msg for phrase in ['how old are you', 'your age', 'favorite color', 'what do you like', 'do you eat', 'are you hungry']):
        intent = 'nlp_gemini_fallback' # Force Gemini for personal/casual curiosity

    # General greeting intent (check last to avoid conflicts)
    elif any(phrase in msg for phrase in ['hi', 'hello', 'hey', 'heya', 'yo', 'hola', 'namaste', 'good morning', 'good afternoon', 'good evening', 'hello there', 'hi bot', 'hi smartedu', 'hello smartedu', 'hey smartedu', 'smartedu are you there', 'gm', 'gn', 'lunch', 'dinner', 'breakfast', 'eat']):
        intent = 'greet'
    elif any(word in msg for word in ['fee', 'fees', 'payment', 'due', 'balance', 'show my fees', 'fee details', 'fee status']):
        intent = 'fee_query'
    elif any(word in msg for word in ['fine', 'fines', 'penalty', 'late fee', 'my fines']):
        intent = 'fine_query'
    elif any(word in msg for word in ['attendance', 'present', 'absent', 'percentage', 'attend', 'show my attendance', 'attendance details']):
        intent = 'attendance_query'
    elif any(phrase in msg for phrase in ['register for', 'enroll in', 'sign up for', 'participate in']):
        intent = 'register_event'
    elif any(word in msg for word in ['event', 'events', 'upcoming', 'schedule', 'seminar', 'workshop', 'fest', 'summit', 'show events', 'list events']):
        intent = 'event_query'
    elif any(word in msg for word in ['announcement', 'announcements', 'notification', 'notifications', 'show announcements']):
        intent = 'announcement_query'
    elif any(word in msg for word in ['admission', 'admit', 'apply', 'application', 'eligibility', 'entrance', 'mht-cet', 'jee', 'seat', 'intake']):
        intent = 'ask_admission'
    elif any(word in msg for word in ['cutoff', 'merit', 'rank', 'percentile']):
        intent = 'ask_cutoff'
    elif any(word in msg for word in ['college', 'campus', 'about', 'info', 'information', 'skn', 'sinhgad', 'approved', 'affiliated', 'established']):
        intent = 'ask_college_info'
    elif any(word in msg for word in ['hostel', 'accommodation', 'boarding', 'room', 'mess', 'lodging']):
        intent = 'ask_hostel'
    elif any(word in msg for word in ['transport', 'bus', 'vehicle', 'commute', 'travel']):
        intent = 'ask_transport'
    elif any(word in msg for word in ['placement', 'recruiter', 'package', 'salary', 'company', 'career']):
        intent = 'ask_placement'
    elif any(word in msg for word in ['scholarship', 'financial aid', 'tfws', 'ebc', 'fee waiver']):
        intent = 'ask_scholarship'
    elif any(word in msg for word in ['document', 'certificate', 'marksheet', 'id proof', 'aadhaar', 'caste']):
        intent = 'ask_documents'
    elif any(word in msg for word in ['guidance', 'suggest', 'which branch', 'best branch', 'change branch']):
        intent = 'ask_guidance'
    elif any(word in msg for word in ['profile', 'who am i', 'student id', 'roll number', 'department', 'class']):
        intent = 'student_info'
    elif any(word in msg for word in ['mark', 'marks', 'score', 'result', 'grade', 'academic', 'performance']):
        intent = 'marks_query'

    if intent == 'unknown':
        from nlp_engine import nlp_engine
        intent, confidence = nlp_engine.predict_intent(msg)
    # If confidence is too low or intent is unknown, fallback to Gemini
    # Check if this is a general conversational query or a college info query
    is_conversational = intent in ['greet', 'goodbye', 'ask_howareyou', 'ask_whoareyou', 'ask_areyoubot', 'ask_creator', 'ask_help']
    is_general_knowledge = intent in ['ask_admission', 'ask_cutoff', 'ask_college_info', 'ask_hostel', 'ask_transport', 'ask_placement', 'ask_scholarship', 'ask_documents', 'ask_facilities', 'ask_library', 'ask_eligibility', 'ask_guidance']
    
    # Special check for general fee structure questions vs individual fee balance
    is_general_fee_query = intent == 'fee_query' and any(word in msg for word in ['structure', 'average', 'category', 'open', 'sc', 'st', 'obc', 'ebc', 'nt', 'sbc', 'tfws', 'girl', 'b.tech', 'm.tech', 'hostel fee', 'mess fee'])

    if confidence < 0.50 or intent == 'nlp_gemini_fallback' or (intent == 'unknown' and not is_conversational) or is_general_knowledge or is_general_fee_query:
        # We need a fallback response for complex questions or static college data
        from datetime import datetime
        now = datetime.now()
        user_context = {
            'role': user_role,
            'name': admin_user.name if admin_user else 'Student',
            'is_admin': user_role in ['admin', 'HOD'],
            'current_time': now.strftime("%I:%M %p"),
            'current_date': now.strftime("%A, %B %d, %Y")
        }
        
        student_db_context = ""
        if user_role not in ['admin', 'HOD']:
            # For students, try to use provided student_id or find it by user_id
            sid = student_id
            if not sid and user_id:
                student_search = Student.query.filter_by(user_id=int(user_id)).first()
                if student_search:
                    sid = student_search.student_id
            
            if sid:
                try:
                    # Fetch basic info, attendance and marks using existing models
                    student = db.session.get(Student, sid)
                    if student:
                        student_class = student.class_info.class_name if student.class_info else "Unknown"
                        student_dept = student.department.department_name if student.department else "Unknown"
                        
                        # Get Attendance
                        attendance_records = Attendance.query.filter_by(student_id=sid).all()
                        total_classes = sum(a.total_lectures or 0 for a in attendance_records)
                        attended_classes = sum(a.attended_lectures or 0 for a in attendance_records)
                        attendance_percent = (attended_classes / total_classes * 100) if total_classes > 0 else 0
                        
                        # Get Marks
                        marks_records = Mark.query.filter_by(student_id=sid).all()
                        marks_info = []
                        for m in marks_records:
                            subj = db.session.get(Subject, m.subject_id)
                            subj_name = subj.subject_name if subj else "Subject"
                            obtained = float(m.obtained_marks or 0)
                            total = float(m.total_marks or 0)
                            marks_info.append(f"{subj_name}: {obtained:.0f}/{total:.0f}")
                        marks_str = ", ".join(marks_info) if marks_info else "No marks recorded"
                        
                        student_db_context = f"Student Name: {student.user.name}, Roll No: {student.roll_no}, Class: {student_class}, Department: {student_dept}. Attendance: {attendance_percent:.1f}% ({attended_classes}/{total_classes} lectures). Marks: {marks_str}."
                except Exception as e:
                    print(f"Error fetching student DB context: {e}")

        gemini_response = nlp_engine.generate_gemini_response(message, intent=intent, user_context=user_context, student_db_context=student_db_context)
        return jsonify({
            'intent': intent if intent != 'unknown' else 'nlp_gemini_response', 
            'response': gemini_response, 
            'data': {},
            'is_ai': True
        })
    
    # Otherwise, continue with structured database routines based on intent


    requested_subject = None
    if intent in ['marks_query', 'attendance_query']:
        try:
            for subject in Subject.query.all():
                s_name = subject.subject_name.lower()
                s_code = subject.subject_code.lower() if subject.subject_code else ''
                m_lower = msg.lower()
                m_norm = m_lower.replace('s ', ' ').replace('s', '')
                s_norm = s_name.replace('s ', ' ').replace('s', '')
                if (s_name in m_lower or s_name + 's' in m_lower or 
                    (len(m_norm) > 4 and m_norm in s_norm) or (len(s_norm) > 4 and s_norm in m_norm) or
                    (s_code and len(s_code) >= 2 and s_code in m_lower)):
                    requested_subject = subject.subject_name
                    break
        except Exception:
            pass

    if any(word in msg for word in ['profile', 'who am i', 'student id', 'roll number', 'department', 'class']):
        return jsonify({'intent': 'student_info', 'response': 'Please visit your Profile page to view student details.', 'data': {}})
    elif any(word in msg for word in ['mark', 'marks', 'score', 'result', 'grade', 'academic', 'performance']):
        # Handle marks query with proper API call
        # Check if admin is asking
        is_admin = user_role == 'admin' or user_role == 'HOD'
        
        if not is_admin:
            # For students, get their student_id
            if not student_id:
                if user_id:
                    student = Student.query.filter_by(user_id=int(user_id)).first()
                    if student:
                        student_id = student.student_id
            if not student_id:
                return jsonify({'intent': 'marks_query', 'response': 'Please log in as a student to view marks details.', 'data': {}}), 200
        

        # Admin handling - show all students or specific student's marks
        if is_admin:
            try:
                import requests
                # Use helper function to detect filters
                filters = detect_admin_filters(msg)
                
                detected_subject = requested_subject
                if detected_subject:
                    # Get all students and filter by subject
                    # For subject queries, get all students regardless of class
                    all_students = Student.query.all()
                    
                    subject_text = f" in {detected_subject.upper()}"
                    response_text = f"📊 **Students Academic Performance{subject_text}**\n"
                    response_text += f"{'='*60}\n\n"
                    
                    for student in all_students:
                        student_marks = Mark.query.filter_by(student_id=student.student_id).all()
                        if student_marks:
                            # Track if we've already shown this student
                            student_shown = False
                            for mark in student_marks:
                                subject = db.session.get(Subject, mark.subject_id)
                                if subject:
                                    # Normalize both detected subject and actual subject name for comparison
                                    detected_normalized = detected_subject.lower().replace('s ', ' ').replace('s', '').strip()
                                    actual_normalized = subject.subject_name.lower().replace('s ', ' ').replace('s', '').strip()
                                    # Check if they match
                                    if detected_normalized == actual_normalized or detected_subject.lower() in subject.subject_name.lower():
                                        if not student_shown:
                                            obtained = float(mark.obtained_marks)
                                            total = float(mark.total_marks)
                                            percentage = (obtained / total * 100) if total > 0 else 0
                                            status_emoji = "✅" if percentage >= 35 else "❌"
                                            response_text += f"📌 **{student.user.name}** (Roll: {student.roll_no})\n"
                                            response_text += f"   └─ {subject.subject_name}: {obtained:.0f}/{total:.0f} ({percentage:.1f}%) {status_emoji}\n\n"
                                            student_shown = True
                                            break
                    
                    return jsonify({'intent': 'marks_query', 'response': response_text, 'data': {}})
                
                if filters['target_student_id']:
                    # Get student object for name
                    student_obj = db.session.get(Student, filters['target_student_id'])
                    student_name = student_obj.user.name if student_obj else 'Student'
                    
                    # Fetch specific student's marks
                    api_response = requests.get(f'http://127.0.0.1:5000/api/students/{filters["target_student_id"]}/marks', timeout=5)
                    if api_response.status_code == 200:
                        marks_data = api_response.json()
                        if marks_data and len(marks_data) > 0:
                            response_text = f"📊 **Academic Performance Report for {student_name}**\n"
                            response_text += f"{'='*60}\n\n"
                            for i, mark in enumerate(marks_data[:10], 1):
                                subject_name = mark.get('subject_name', f'Subject {i}')
                                obtained = float(mark.get('obtained_marks', 0))
                                total = float(mark.get('total_marks', 100))
                                percentage = (obtained / total * 100) if total > 0 else 0
                                status_emoji = "✅" if percentage >= 35 else "❌"
                                status_text = "PASS" if percentage >= 35 else "FAIL"
                                response_text += f"📌 **{subject_name}**\n"
                                response_text += f"   └─ Marks: {obtained:.0f}/{total:.0f}  |  Percentage: {percentage:.1f}%  |  Status: {status_emoji} {status_text}\n\n"
                            response_text += f"{'='*60}"
                            return jsonify({'intent': 'marks_query', 'response': response_text, 'data': {'marks': marks_data}})
                        else:
                            return jsonify({'intent': 'marks_query', 'response': f'✅ **Good news!** There are no pending marks for {student_name}.', 'data': {}})
                    else:
                        return jsonify({'intent': 'marks_query', 'response': 'Unable to fetch marks details.', 'data': {}})
                else:
                    # Admin wants all students' marks - query directly from database
                    # Apply class filter if specified
                    query = Student.query
                    
                    if filters['class_filter']:
                        # Filter by class name (search in class.class_name)
                        query = query.join(Class).filter(Class.class_name.ilike(f'%{filters["class_filter"]}%'))
                    
                    all_students = query.all()
                    
                    marks_found = False
                    class_text = f" ({filters['class_filter']})" if filters['class_filter'] else ""
                    response_text = f"📊 **Students Academic Performance{class_text}**\n"
                    response_text += f"{'='*60}\n\n"
                    
                    for student in all_students[:20]:  # Show max 20 students
                            # JOIN with Subject to get accurate totals for the summary
                            marks_and_subjects = db.session.query(Mark, Subject).join(
                                Subject, Mark.subject_id == Subject.subject_id
                            ).filter(Mark.student_id == student.student_id).all()
                            
                            class_total_marks = sum(float(s.total_marks or 0) for m, s in marks_and_subjects)
                            class_total_obtained = sum(float(m.obtained_marks or 0) for m, s in marks_and_subjects)
                            
                            if class_total_marks > 0:
                                marks_found = True
                                percentage = (class_total_obtained / class_total_marks * 100)
                                response_text += f"📌 **{student.user.name}** (Roll: {student.roll_no})\n"
                                response_text += f"   └─ Total: {class_total_obtained:.0f}/{class_total_marks:.0f} ({percentage:.1f}%)\n\n"
                    
                    if not marks_found:
                        response_text = f"✅ **Good news!** There are no pending marks for students{class_text}."
                    else:
                        response_text += f"{'='*60}"
                        
                    return jsonify({'intent': 'marks_query', 'response': response_text, 'data': {}})
            except Exception as e:
                print(f"Error in admin marks query: {e}")
                return jsonify({'intent': 'marks_query', 'response': 'Unable to fetch marks details.', 'data': {}})
        
        # Student handling - use the same API endpoint that Student Services uses
        try:
            import requests
            api_response = requests.get(f'http://127.0.0.1:5000/api/students/{student_id}/marks', timeout=5)
            if api_response.status_code == 200:
                marks_data = api_response.json()
                if marks_data and len(marks_data) > 0:
                    # If specific subject requested, filter results
                    if requested_subject:
                        matched_marks = []
                        seen_subjects = set()  # Track unique subjects to avoid duplicates
                        for mark in marks_data:
                            mark_subject = mark.get('subject_name', '').lower()
                            # Check if this matches our requested subject
                            matches = requested_subject.lower() in mark_subject.lower() or requested_subject.lower().replace('s', '') in mark_subject.lower().replace('s', '')
                            
                            # Only add if it matches AND we haven't seen this exact subject name before
                            if matches and mark_subject not in seen_subjects:
                                matched_marks.append(mark)
                                seen_subjects.add(mark_subject)
                        
                        if matched_marks:
                            response_text = f"📊 **Subject: {requested_subject.title()}**\n"
                            response_text += f"{'='*50}\n\n"
                            for i, mark in enumerate(matched_marks, 1):
                                subject_name = mark.get('subject_name', 'N/A')
                                obtained = float(mark.get('obtained_marks', 0))
                                total = float(mark.get('total_marks', 100))
                                percentage = (obtained / total * 100) if total > 0 else 0
                                status_emoji = "✅" if percentage >= 35 else "❌"
                                status_text = "PASS" if percentage >= 35 else "FAIL"
                                exam_date = mark.get('exam_date', 'N/A')
                                
                                response_text += f"📌 **Subject:**     {subject_name}\n"
                                response_text += f"🎯 **Marks:**       {obtained:.0f}/{total:.0f} ({percentage:.1f}%)\n"
                                response_text += f"📅 **Exam Date:**   {exam_date}\n"
                                response_text += f"🔖 **Status:**      {status_emoji} {status_text}\n"
                                if i < len(matched_marks):
                                    response_text += f"\n{'-'*50}\n\n"
                            
                            response_text += f"{'='*50}"
                            return jsonify({'intent': 'marks_query', 'response': response_text, 'data': {'marks': matched_marks}})
                        else:
                            return jsonify({'intent': 'marks_query', 'response': f'✅ **Good news!** There are no pending marks for {requested_subject}.', 'data': {}})
                    else:
                        # Show all subjects with better formatting
                        response_text = f"📊 **Academic Performance Report**\n"
                        response_text += f"{'='*60}\n\n"
                        # Use JOIN to ensure individual totals are correct
                        student_marks_with_subjects = db.session.query(Mark, Subject).join(
                            Subject, Mark.subject_id == Subject.subject_id
                        ).filter(Mark.student_id == student_id).all()
                        
                        total_marks = 0
                        total_obtained = 0
                        
                        # Deduplicate subjects if necessary (though query should be clean)
                        seen_subjects = set()
                        for mark, subject in student_marks_with_subjects:
                            subject_name = subject.subject_name
                            if subject_name not in seen_subjects:
                                seen_subjects.add(subject_name)
                                
                                obtained = float(mark.obtained_marks or 0)
                                total = float(subject.total_marks or 100)
                                percentage = (obtained / total * 100) if total > 0 else 0
                                status_emoji = "✅" if percentage >= 35 else "❌"
                                status_text = "PASS" if percentage >= 35 else "FAIL"
                                
                                response_text += f"📌 **{subject_name}**\n"
                                response_text += f"   └─ Marks: {obtained:.0f}/{total:.0f}  |  Percentage: {percentage:.1f}%  |  Status: {status_emoji} {status_text}\n\n"
                                
                                total_obtained += obtained
                                total_marks += total
                        
                        if total_marks > 0:
                            overall_percentage = (total_obtained / total_marks * 100)
                            response_text += f"{'='*60}\n"
                            response_text += f"📈 **Overall Performance:** {total_obtained:.1f}/{total_marks:.1f} ({overall_percentage:.1f}%)\n"
                            response_text += f"{'='*60}"
                        
                        return jsonify({'intent': 'marks_query', 'response': response_text, 'data': {'marks': marks_data}})
                else:
                    return jsonify({'intent': 'marks_query', 'response': '✅ **Good news!** There are no pending marks for your account.', 'data': {}})
            else:
                return jsonify({'intent': 'marks_query', 'response': 'Unable to fetch marks details at the moment.', 'data': {}})
        except Exception as e:
            return jsonify({'intent': 'marks_query', 'response': 'Unable to fetch marks details at the moment.', 'data': {}})
    elif any(word in msg for word in ['help', 'what can you do', 'features', 'how to use', 'assistance']):
        return jsonify({'intent': 'help_query', 'response': 'I can help with fees, attendance, and events. For other features, please explore the dashboard.', 'data': {}})
    
    data = {}
    if intent == 'greet':
        # Detect greeting type and respond accordingly
        greeting_keywords = {
            'simple': ['hi', 'hello', 'hey', 'heya', 'yo', 'hola', 'namaste'],
            'polite': ['good morning', 'good afternoon', 'good evening'],
            'friendly': ['what\'s up', 'how\'s it going', 'how are you', 'how are you doing', 'how\'s your day', 'how are things', 'what\'s new'],
            'introductory': ['who are you', 'what\'s your name', 'are you a bot', 'are you real', 'are you human', 'who made you'],
            'return': ['nice to meet you', 'glad to see you', 'good to see you again', 'long time no see'],
            'checking': ['are you there', 'you there', 'can you help me', 'anyone here', 'are you online']
        }
        
        msg_lower = msg.lower()
        greeting_type = None
        
        # Detect greeting type
        if any(word in msg_lower for word in greeting_keywords['polite']):
            greeting_type = 'polite'
        elif any(word in msg_lower for word in greeting_keywords['friendly']):
            greeting_type = 'friendly'
        elif any(word in msg_lower for word in greeting_keywords['introductory']):
            greeting_type = 'introductory'
        elif any(word in msg_lower for word in greeting_keywords['return']):
            greeting_type = 'return'
        elif any(word in msg_lower for word in greeting_keywords['checking']):
            greeting_type = 'checking'
        else:
            greeting_type = 'simple'
        
        # Personalized responses based on greeting type
        if greeting_type == 'polite':
            time_of_day = datetime.now().hour
            if time_of_day < 12:
                response_text = "Good morning! ≡ƒîà I'm SmartEdu, your friendly college assistant. I'm here to help you with admission information, fees, hostels, placements, and more. How can I assist you today?"
            elif time_of_day < 17:
                response_text = "Good afternoon! ΓÿÇ∩╕Å I'm SmartEdu, your college information assistant. Ask me about SKN Sinhgad College's admissions, courses, fees, or anything else you'd like to know!"
            else:
                response_text = "Good evening! ≡ƒîÖ I'm SmartEdu, here to help you with college information. I can provide details about admissions, hostels, placements, and more. What would you like to know?"
        elif greeting_type == 'friendly':
            if any(word in msg_lower for word in ['how are you', 'how are you doing']):
                response_text = "I'm doing great, thanks for asking! ≡ƒÿè I'm SmartEdu, your college assistant. I'm here and ready to help you with admission information for SKN Sinhgad College. What can I help you with today?"
            else:
                response_text = "Hey there! ≡ƒæï I'm SmartEdu, your college information buddy! I'm doing great and excited to help you. Want to know about admissions, fees, hostels, or placements at SKN Sinhgad College?"
        elif greeting_type == 'introductory':
            if 'name' in msg_lower:
                response_text = "Hi! My name is SmartEdu ≡ƒñû - your smart college assistant! I was created to help students and visitors get information about SKN Sinhgad College of Engineering. I can answer questions about admissions, courses, fees, hostels, and more. Nice to meet you!"
            elif 'bot' in msg_lower or 'human' in msg_lower or 'real' in msg_lower:
                response_text = "Yes, I'm SmartEdu! ≡ƒñû I'm an AI chatbot created to help you with college information. While I'm not human, I'm here 24/7 to assist you with everything about SKN Sinhgad College - admissions, courses, fees, placements, and more. How can I help you?"
            else:
                response_text = "I'm SmartEdu! ≡ƒñû Your smart college information assistant at SKN Sinhgad College. I'm designed to help you with admissions, courses, fees, hostels, placements, and more. I'm here to make your college journey easier!"
        elif greeting_type == 'return':
            response_text = "Welcome back! ≡ƒÖî Good to see you again! I'm SmartEdu, and I'm always here to help you with college information. Ready to assist you with admissions, fees, or anything else you need. What would you like to know?"
        elif greeting_type == 'checking':
            response_text = "Yes, I'm here! ≡ƒæï Hello! I'm SmartEdu, your college assistant, and I'm online and ready to help. I can assist you with admission information, fees, hostels, placements, or any questions about SKN Sinhgad College. How can I help you today?"
        else:  # simple greeting
            response_text = "Hello! ≡ƒæï I'm SmartEdu, your college information assistant. Nice to meet you! I can help you with admissions, courses, fees, hostels, placements, and more at SKN Sinhgad College. What would you like to know?"
    
    elif intent == 'goodbye':
        import random
        goodbye_responses = [
            "Goodbye ≡ƒæï! Have a great day ahead!",
            "Bye for now! ≡ƒÿè Don't forget to study smart!",
            "See you soon! ≡ƒôÜ SmartEdu is always here when you need help.",
            "Take care! ≡ƒîƒ Come back anytime for more info.",
            "Goodbye from SmartEdu ≡ƒñû ΓÇö wishing you success ahead!"
        ]
        response_text = random.choice(goodbye_responses)
    
    elif intent == 'ask_howareyou':
        import random
        howareyou_responses = [
            "I'm doing great! ≡ƒñû Ready to help you with college info. How about you?",
            "Feeling smart as always ≡ƒÿÄ! What brings you here today?",
            "All systems go! ≡ƒÜÇ How can I assist you today?",
            "I'm always good when students come to chat with me! ≡ƒÿè"
        ]
        response_text = random.choice(howareyou_responses)
    
    elif intent == 'ask_whoareyou':
        import random
        whoareyou_responses = [
            "I'm **SmartEdu**, your AI-powered college assistant chatbot. I can help you with admissions, fees, results, and more!",
            "I'm SmartEdu ≡ƒñû ΓÇö your digital academic buddy built to guide students and answer queries.",
            "I'm SmartEdu, developed to make your campus life easier and smarter! ≡ƒÄô"
        ]
        response_text = random.choice(whoareyou_responses)
    
    elif intent == 'ask_areyoubot':
        import random
        areyoubot_responses = [
            "Yes, I'm a chatbot ≡ƒñû ΓÇö built smart to help students like you!",
            "Absolutely! I'm a virtual assistant designed to answer all your college-related queries.",
            "Yup! I'm your friendly digital assistant ΓÇö SmartEdu at your service. ≡ƒÆ¼"
        ]
        response_text = random.choice(areyoubot_responses)
    
    elif intent == 'ask_creator':
        import random
        creator_responses = [
            "I was created by the SmartEdu development team ≡ƒÆ╗ ΓÇö guided by talented computer science students!",
            "SmartEdu was developed by passionate tech minds from your college. ≡ƒÄô",
            "I was built with love Γ¥ñ∩╕Å and code ≡ƒÆ╗ by the SmartEdu team!"
        ]
        response_text = random.choice(creator_responses)
    
    elif intent == 'ask_help':
        import random
        help_responses = [
            "Sure! ≡ƒÿè I can help you with admissions, fees, courses, results, events, and more. What do you want to know?",
            "I'm here to assist! ≡ƒÄô You can ask about academics, hostel, placement, or any campus info.",
            "Of course! ≡ƒñû Tell me your question ΓÇö admissions, attendance, marks, or events?",
            "Always ready to help you! ≡ƒÆ¼ What's your query today?"
        ]
        response_text = random.choice(help_responses)
    
    elif intent == 'fee_query':
        # Check if admin is asking
        is_admin_fee = user_role == 'admin' or user_role == 'HOD'
        
        if not is_admin_fee:
            # For students, get their student_id
            if not student_id:
                # fallback: derive student_id from user_id if provided
                if user_id:
                    student = Student.query.filter_by(user_id=int(user_id)).first()
                    if student:
                        student_id = student.student_id
            if not student_id:
                return jsonify({'intent': 'fee_query', 'response': 'Please log in as a student to view fee details.', 'data': {}}), 200
        
        # Admin handling for fees
        if is_admin_fee:
            try:
                import requests
                # Use helper function to detect filters
                filters = detect_admin_filters(msg)
                
                if filters['target_student_id']:
                    # Get student object for name
                    student_obj = db.session.get(Student, filters['target_student_id'])
                    student_name = student_obj.user.name if student_obj else 'Student'
                    
                    # Fetch specific student's fees
                    api_response = requests.get(f'http://127.0.0.1:5000/api/students/{filters["target_student_id"]}/fees', timeout=5)
                    if api_response.status_code == 200:
                        fee_data = api_response.json()
                        if fee_data and 'total_amount' in fee_data:
                            total = fee_data.get('total_amount', 0)
                            paid = fee_data.get('paid_amount', 0)
                            due = fee_data.get('due_amount', 0)
                            
                            response_text = f"💳 **Fee Payment Details for {student_name}**\n"
                            response_text += f"{'='*50}\n"
                            response_text += f"\n📌 **Total Fee:**      ₹{total:,.2f}\n"
                            response_text += f"✅ **Paid Amount:**   ₹{paid:,.2f}\n"
                            response_text += f"⏳ **Due Amount:**    ₹{due:,.2f}\n"
                            response_text += f"\n📋 **Status:**         {fee_data.get('payment_status', 'N/A')}"
                            if fee_data.get('last_payment_date'):
                                response_text += f"\n📅 **Last Payment:**   {fee_data.get('last_payment_date')}"
                            response_text += f"\n{'='*50}"
                            return jsonify({'intent': 'fee_query', 'response': response_text, 'data': fee_data})
                        else:
                            return jsonify({'intent': 'fee_query', 'response': f'✅ **Good news!** There are no pending fees for {student_name}.', 'data': {}})
                else:
                    # Admin wants all students' fees - query directly from database
                    # Apply class filter if specified
                    query = Student.query
                    
                    if filters['class_filter']:
                        query = query.join(Class).filter(Class.class_name.ilike(f'%{filters["class_filter"]}%'))
                    
                    all_students = query.all()
                    
                    fines_found = False
                    class_text = f" ({filters['class_filter']})" if filters['class_filter'] else ""
                    response_text = f"💳 **Students Fee Status{class_text}**\n"
                    response_text += f"{'='*60}\n\n"
                    
                    for student in all_students[:20]:  # Show max 20 students
                        student_fees = Fee.query.filter_by(student_id=student.student_id).first()
                        if student_fees:
                            fines_found = True
                            fee_data = student_fees.to_dict()
                            status_emoji = "✅" if fee_data.get('due_amount', 0) == 0 else "⚠️"
                            response_text += f"📌 **{student.user.name}** (Roll: {student.roll_no})\n"
                            response_text += f"   └─ Paid: ₹{fee_data.get('paid_amount', 0):,.0f} / Total: ₹{fee_data.get('total_amount', 0):,.0f}  |  Status: {status_emoji} {fee_data.get('payment_status', 'N/A')}\n\n"
                    
                    if not fines_found:
                        response_text = f"✅ **Good news!** There are no pending fees for students{class_text}."
                    else:
                        response_text += f"{'='*60}"
                        
                    return jsonify({'intent': 'fee_query', 'response': response_text, 'data': {}})
            except Exception as e:
                print(f"Error in admin fee query: {e}")
                return jsonify({'intent': 'fee_query', 'response': 'Unable to fetch fee details.', 'data': {}})
        
        # Student handling
        try:
            student_fees = Fee.query.filter_by(student_id=student_id).first()
            if student_fees:
                fee_data = student_fees.to_dict()
                total = fee_data.get('total_amount', 0)
                paid = fee_data.get('paid_amount', 0)
                due = fee_data.get('due_amount', 0)
                
                response_text = f"💳 **Fee Payment Details**\n"
                response_text += f"{'='*50}\n"
                response_text += f"\n📌 **Total Fee:**      ₹{total:,.2f}\n"
                response_text += f"✅ **Paid Amount:**   ₹{paid:,.2f}\n"
                response_text += f"⏳ **Due Amount:**    ₹{due:,.2f}\n"
                response_text += f"\n📋 **Status:**         {fee_data.get('payment_status', 'N/A')}"
                if fee_data.get('last_payment_date'):
                    response_text += f"\n📅 **Last Payment:**   {fee_data.get('last_payment_date')}"
                response_text += f"\n{'='*50}"
                data = fee_data
            else:
                response_text = "✅ **Good news!** There are no pending fees for your account."
                data = {}
        except Exception as e:
            print(f"Error fetching fee details for student: {e}")
            response_text = "Unable to fetch fee details at the moment."
            data = {}
            
    elif intent == 'fine_query':
        is_admin_fine = user_role == 'admin' or user_role == 'HOD'
        
        if not is_admin_fine:
            if not student_id:
                if user_id:
                    student = Student.query.filter_by(user_id=int(user_id)).first()
                    if student:
                        student_id = student.student_id
            if not student_id:
                return jsonify({'intent': 'fine_query', 'response': 'Please log in as a student to view fine details.', 'data': {}}), 200
        
        # Admin handling for fines
        if is_admin_fine:
            try:
                filters = detect_admin_filters(msg)
                
                if filters['target_student_id']:
                    student_obj = db.session.get(Student, filters['target_student_id'])
                    student_name = student_obj.user.name if student_obj else 'Student'
                    
                    student_fines = Fine.query.filter_by(student_id=filters['target_student_id']).all()
                    if student_fines:
                        response_text = f"⚠️ **Fines and Penalties for {student_name}**\n"
                        response_text += f"{'='*50}\n\n"
                        total_fines = 0
                        for fine in student_fines:
                            fine_data = fine.to_dict()
                            amount = float(fine_data.get('amount', 0))
                            total_fines += amount
                            status = "⌛ " + fine_data.get('status', 'Pending') if fine_data.get('status') != 'Approved' else "✅ Cleared"
                            response_text += f"📌 **{fine_data.get('reason', 'Penalty')}**\n"
                            response_text += f"   └─ Amount: ₹{amount:,.2f}  |  Status: {status}\n\n"
                        
                        response_text += f"{'='*50}\n"
                        response_text += f"**Total Fines:** ₹{total_fines:,.2f}\n"
                        return jsonify({'intent': 'fine_query', 'response': response_text, 'data': {'fines': [f.to_dict() for f in student_fines]}})
                    else:
                        return jsonify({'intent': 'fine_query', 'response': f"✅ **Good news!** There are no pending fines for {student_name}.", 'data': {}})
                else:
                    query = Student.query
                    if filters['class_filter']:
                        query = query.join(Class).filter(Class.class_name.ilike(f'%{filters["class_filter"]}%'))
                    
                    all_students = query.all()
                    class_text = f" ({filters['class_filter']})" if filters['class_filter'] else ""
                    
                    fines_found = False
                    response_text = f"⚠️ **Students Fines Status{class_text}**\n"
                    response_text += f"{'='*60}\n\n"
                    
                    for student in all_students:
                        student_fines = Fine.query.filter_by(student_id=student.student_id).all()
                        if student_fines:
                            fines_found = True
                            student_total = sum(float(f.amount or 0) for f in student_fines)
                            response_text += f"📌 **{student.user.name}** (Roll: {student.roll_no})\n"
                            response_text += f"   └─ Total Fines: ₹{student_total:,.2f} across {len(student_fines)} record(s)\n\n"
                    
                    if not fines_found:
                        response_text = f"✅ **Good news!** There are no pending fines for students{class_text}."
                    else:
                        response_text += f"{'='*60}"
                        
                    return jsonify({'intent': 'fine_query', 'response': response_text, 'data': {}})
            except Exception as e:
                print(f"Error in admin fine query: {e}")
                return jsonify({'intent': 'fine_query', 'response': 'Unable to fetch fine details.', 'data': {}})
                
        # Student handling
        try:
            student_fines = Fine.query.filter_by(student_id=student_id).all()
            if student_fines:
                response_text = f"⚠️ **My Fines and Penalties**\n"
                response_text += f"{'='*50}\n\n"
                total_fines = 0
                for fine in student_fines:
                    fine_data = fine.to_dict()
                    amount = float(fine_data.get('amount', 0))
                    total_fines += amount
                    status = "⌛ " + fine_data.get('status', 'Pending') if fine_data.get('status') != 'Approved' else "✅ Cleared"
                    response_text += f"📌 **{fine_data.get('reason', 'Penalty')}**\n"
                    response_text += f"   └─ Amount: ₹{amount:,.2f}  |  Status: {status}\n\n"
                
                response_text += f"{'='*50}\n"
                response_text += f"**Total Fines:** ₹{total_fines:,.2f}\n"
                data = {'fines': [f.to_dict() for f in student_fines]}
            else:
                response_text = "✅ **Good news!** There are no pending fines for your account."
                data = {}
        except Exception as e:
            print(f"Error fetching fine details for student: {e}")
            response_text = "Unable to fetch fine details at the moment."
            data = {}
            
    elif intent == 'attendance_query':
        # Check if admin is asking
        is_admin_attendance = user_role == 'admin' or user_role == 'HOD'
        
        if not is_admin_attendance:
            if not student_id:
                user_id = payload.get('user_id')
                if user_id:
                    student = Student.query.filter_by(user_id=int(user_id)).first()
                    if student:
                        student_id = student.student_id
            if not student_id:
                return jsonify({'intent': 'attendance_query', 'response': 'Please log in as a student to view attendance details.', 'data': {}}), 200
        
        msg = payload.get('message', '').lower()
        
        # Admin handling for attendance
        if is_admin_attendance:
            try:
                import requests
                # Use helper function to detect filters
                filters = detect_admin_filters(msg)
                
                if filters['target_student_id']:
                    # Get student object for name
                    student_obj = db.session.get(Student, filters['target_student_id'])
                    student_name = student_obj.user.name if student_obj else 'Student'
                    
                    # Fetch specific student's attendance
                    api_response = requests.get(f'http://127.0.0.1:5000/api/students/{filters["target_student_id"]}/attendance-summary', timeout=5)
                    if api_response.status_code == 200:
                        attendance_data = api_response.json()
                        if attendance_data and len(attendance_data) > 0:
                            total_present = sum(record.get('present_count', 0) for record in attendance_data)
                            total_classes = sum(record.get('total_classes', 0) for record in attendance_data)
                            attendance_percentage = (total_present / total_classes * 100) if total_classes > 0 else 0
                            
                            response_text = f"📈 **Attendance for {student_name}**\n"
                            response_text += f"{'='*50}\n\n"
                            response_text += f"📊 **Overall:** {attendance_percentage:.1f}% ({total_present}/{total_classes})\n\n"
                            
                            for record in attendance_data[:5]:
                                subject_name = record.get('subject_name', 'N/A')
                                attended = record.get('present_count', 0)
                                total = record.get('total_classes', 0)
                                percentage = (attended / total * 100) if total > 0 else 0
                                response_text += f"📌 **{subject_name}:** {attended}/{total} ({percentage:.1f}%)\n"
                            
                            response_text += f"\n{'='*50}"
                            return jsonify({'intent': 'attendance_query', 'response': response_text, 'data': {'attendance': attendance_data}})
                        else:
                            return jsonify({'intent': 'attendance_query', 'response': f'✅ **Good news!** There are no pending attendance records for {student_name}.', 'data': {}})
                else:
                    # Admin wants all students' attendance summary - query directly from database
                    # Apply class filter if specified
                    query = Student.query
                    
                    if filters['class_filter']:
                        query = query.join(Class).filter(Class.class_name.ilike(f'%{filters["class_filter"]}%'))
                    
                    all_students = query.all()
                    
                    attendance_found = False
                    class_text = f" ({filters['class_filter']})" if filters['class_filter'] else ""
                    response_text = f"📈 **Students Attendance Summary{class_text}**\n"
                    response_text += f"{'='*60}\n\n"
                    
                    for student in all_students[:20]:  # Show max 20 students
                        attendance_records = Attendance.query.filter_by(student_id=student.student_id).all()
                        if attendance_records:
                            attendance_found = True
                            total_present = sum(record.present_count for record in attendance_records)
                            total_classes = sum(record.total_classes for record in attendance_records)
                            percentage = (total_present / total_classes * 100) if total_classes > 0 else 0
                            status_emoji = "✅" if percentage >= 75 else "⚠️" if percentage >= 60 else "❌"
                            response_text += f"📌 **{student.user.name}** (Roll: {student.roll_no})\n"
                            response_text += f"   └─ Attendance: {total_present}/{total_classes} ({percentage:.1f}%) {status_emoji}\n\n"
                    
                    if not attendance_found:
                        response_text = f"✅ **Good news!** There are no pending attendance records for students{class_text}."
                    else:
                        response_text += f"{'='*60}"
                        
                    return jsonify({'intent': 'attendance_query', 'response': response_text, 'data': {}})
            except Exception as e:
                print(f"Error in admin attendance query: {e}")
                return jsonify({'intent': 'attendance_query', 'response': 'Unable to fetch attendance details.', 'data': {}})
        
        # Use the same API endpoint that Student Services uses
        try:
            import requests
            api_response = requests.get(f'http://127.0.0.1:5000/api/students/{student_id}/attendance-summary', timeout=5)
            if api_response.status_code == 200:
                attendance_data = api_response.json()
                if attendance_data and len(attendance_data) > 0:
                    # If specific subject requested, filter results
                    if requested_subject:
                        matched_attendance = []
                        seen_subjects = set()  # Track unique subjects to avoid duplicates
                        for record in attendance_data:
                            subject_name = record.get('subject_name', '').lower()
                            # Check if this matches our requested subject
                            matches = requested_subject.lower() in subject_name.lower() or requested_subject.lower().replace('s', '') in subject_name.lower().replace('s', '')
                            
                            # Only add if it matches AND we haven't seen this exact subject name before
                            if matches and subject_name not in seen_subjects:
                                matched_attendance.append(record)
                                seen_subjects.add(subject_name)
                        
                        if matched_attendance:
                            response_text = f"📈 **Attendance: {requested_subject.title()}**\n"
                            response_text += f"{'='*50}\n\n"
                            for record in matched_attendance:
                                subject_name = record.get('subject_name', 'N/A')
                                attended = record.get('present_count', 0)
                                absent = record.get('absent_count', 0)
                                total = record.get('total_classes', 0)
                                percentage = (attended / total * 100) if total > 0 else 0
                                status_emoji = "✅" if percentage >= 75 else "⚠️" if percentage >= 60 else "❌"
                                status_text = "Good" if percentage >= 75 else "Moderate" if percentage >= 60 else "Low"
                                
                                response_text += f"📌 **Subject:**     {subject_name}\n"
                                response_text += f"📊 **Classes:**     {total} (Present: {attended} | Absent: {absent})\n"
                                response_text += f"📈 **Percentage:**  {percentage:.1f}%\n"
                                response_text += f"🔖 **Status:**      {status_emoji} {status_text} Attendance\n"
                            
                            response_text += f"\n{'='*50}"
                            data = {'attendance': matched_attendance}
                        else:
                            response_text = f"✅ **Good news!** There are no pending attendance records for {requested_subject}."
                            data = {}
                    else:
                        # Calculate overall attendance from summary data
                        total_present = sum(record.get('present_count', 0) for record in attendance_data)
                        total_classes = sum(record.get('total_classes', 0) for record in attendance_data)
                        attendance_percentage = (total_present / total_classes * 100) if total_classes > 0 else 0
                        
                        # Format as clean report
                        response_text = f"📈 **Attendance Summary Report**\n"
                        response_text += f"{'='*60}\n\n"
                        response_text += f"📊 **Overall Statistics:**\n"
                        response_text += f"   └─ Total Classes:   {total_classes}\n"
                        response_text += f"   └─ Present:         {total_present}\n"
                        response_text += f"   └─ Absent:          {total_classes - total_present}\n"
                        response_text += f"   └─ Overall:          {attendance_percentage:.1f}%\n\n"
                        
                        # Add subject-wise breakdown (deduplicate subjects)
                        response_text += f"📚 **Subject-wise Attendance:**\n\n"
                        seen_subjects = set()
                        unique_attendance = []
                        
                        # Deduplicate subjects - keep only first occurrence
                        for record in attendance_data:
                            subject_name = record.get('subject_name', 'Unknown').lower()
                            if subject_name not in seen_subjects:
                                seen_subjects.add(subject_name)
                                unique_attendance.append(record)
                        
                        # Display unique subjects
                        for i, record in enumerate(unique_attendance[:8], 1):  # Show max 8 subjects
                            subject_name = record.get('subject_name', f'Subject {i}')
                            attended = record.get('present_count', 0)
                            total = record.get('total_classes', 0)
                            percentage = (attended / total * 100) if total > 0 else 0
                            status_emoji = "✅" if percentage >= 75 else "⚠️" if percentage >= 60 else "❌"
                            status_text = "Good" if percentage >= 75 else "Moderate" if percentage >= 60 else "Low"
                            
                            response_text += f"📌 **{subject_name}**\n"
                            response_text += f"   └─ {attended}/{total} ({percentage:.1f}%)  |  Status: {status_emoji} {status_text}\n\n"
                        
                        if len(unique_attendance) > 8:
                            response_text += f"... and {len(unique_attendance) - 8} more subjects\n\n"
                        
                        # Overall status
                        status_emoji = "✅" if attendance_percentage >= 75 else "⚠️" if attendance_percentage >= 60 else "❌"
                        status_text = "Good" if attendance_percentage >= 75 else "Moderate" if attendance_percentage >= 60 else "Low - Please improve"
                        response_text += f"{'='*60}\n"
                        response_text += f"🔖 **Overall Status:** {status_emoji} {status_text} Attendance\n"
                        response_text += f"{'='*60}"
                        
                        data = {'attendance_summary': attendance_data, 'overall_percentage': attendance_percentage}
                else:
                    response_text = "✅ **Good news!** There are no pending attendance records for your account."
                    data = {}
            else:
                response_text = "Unable to fetch attendance details at the moment."
                data = {}
        except Exception as e:
            response_text = "Unable to fetch attendance details at the moment."
            data = {}
    elif intent == 'event_query':
        events = get_upcoming_events()
        data = {'events': events}
        if events:
            response_text = f"📅 **Upcoming Events ({len(events)} total)**\n"
            response_text += f"{'='*60}\n\n"
            
            # Event type information
            event_type_info = {
                'workshop': {'icon': '🔧', 'form': 'Workshop/Seminar Registration'},
                'seminar': {'icon': '🎓', 'form': 'Workshop/Seminar Registration'},
                'hackathon': {'icon': '💻', 'form': 'Hackathon Event Registration'},
                'club_event': {'icon': '🎭', 'form': 'Club Event Registration'},
                'competition': {'icon': '🏆', 'form': 'Competition Registration'},
                'conference': {'icon': '🏛️', 'form': 'College Event Registration'},
                'cultural': {'icon': '🎨', 'form': 'College Event Registration'},
                'sports': {'icon': '⚽', 'form': 'College Event Registration'},
                'academic': {'icon': '📚', 'form': 'College Event Registration'},
                'general': {'icon': '📅', 'form': 'College Event Registration'}
            }
            
            for i, event in enumerate(events[:8], 1):  # Show max 8 events
                title = event.get('title', 'Untitled Event')
                date = event.get('event_date', 'TBD')
                time = event.get('event_time', 'TBD')
                location = event.get('location', 'TBD')
                event_type = event.get('event_type', 'general')
                type_info = event_type_info.get(event_type, event_type_info['general'])
                desc = event.get('description', '')
                if len(desc) > 100:
                    desc = desc[:100] + '...'
                
                response_text += f"📍 **{title}**\n"
                response_text += f"   └─ {type_info['icon']} Type:       {event_type.title()}\n"
                response_text += f"   └─ 📅 Date:       {date}\n"
                response_text += f"   └─ 🕐 Time:       {time}\n"
                response_text += f"   └─ 📍 Location:   {location}\n"
                response_text += f"   └─ 📝 Form:       {type_info['form']}\n"
                if desc:
                    response_text += f"   └─ 📄 About:      {desc}\n"
                response_text += "\n"
            
            if len(events) > 8:
                response_text += f"... and {len(events) - 8} more events\n"
            
            response_text += f"{'='*60}\n"
            response_text += f"💡 **Registration:** Each event uses a specialized Google Form for professional registration.\n"
            response_text += f"🔗 **Forms Available:** Workshop, Hackathon, Club Events, Competitions, and General Events.\n"
            response_text += f"{'='*60}"
        else:
            response_text = "✅ **Good news!** There are no upcoming events."
    elif intent == 'register_event':
        if not student_id and user_id:
            student = Student.query.filter_by(user_id=int(user_id)).first()
            if student:
                student_id = student.student_id
        
        if not student_id:
            return jsonify({'intent': 'register_event', 'response': 'Please log in as a student to register for events.', 'data': {}}), 200

        # Extract event name from message
        msg_clean = msg.lower()
        for phrase in ['register for', 'enroll in', 'sign up for', 'participate in']:
            msg_clean = msg_clean.replace(phrase, '')
        msg_clean = msg_clean.strip()
            
        events = Event.query.filter_by(is_active=True).all()
        matched_event = None
        if msg_clean:
            for ev in events:
                if msg_clean in ev.title.lower() or ev.title.lower() in msg_clean:
                    matched_event = ev
                    break

        if not matched_event:
            response_text = "Which event would you like to register for? Please specify the exact event name. Say 'show events' to see the list."
            data = {}
        else:
            existing_reg = EventRegistration.query.filter_by(event_id=matched_event.event_id, student_id=student_id).first()
            if existing_reg:
                response_text = f"You are already registered for **{matched_event.title}**! 📅"
                data = {'event': matched_event.to_dict()}
            else:
                try:
                    new_reg = EventRegistration(event_id=matched_event.event_id, student_id=student_id, registered_at=datetime.utcnow())
                    db.session.add(new_reg)
                    
                    if matched_event.current_participants is not None:
                        matched_event.current_participants += 1
                        
                    db.session.commit()
                    response_text = f"✅ **Success!** You have successfully registered for **{matched_event.title}**. We have sent a calendar invite to your email."
                    data = {'event': matched_event.to_dict(), 'registration_success': True}
                except Exception as e:
                    db.session.rollback()
                    print(f"Event registration error: {e}")
                    response_text = f"Sorry, there was an error registering for **{matched_event.title}**. Please try again later."
                    data = {}
    elif intent == 'announcement_query':
        # Get announcements from notifications table
        try:
            user_id = payload.get('user_id')
            if not user_id:
                return jsonify({'intent': 'announcement_query', 'response': 'Please log in to view announcements.', 'data': {}})
            
            import requests
            api_response = requests.get(f'http://127.0.0.1:5000/api/users/{user_id}/notifications', timeout=5)
            if api_response.status_code == 200:
                notifications = api_response.json()
                if notifications and len(notifications) > 0:
                    response_text = f"📢 **Announcements ({len(notifications)} total)**\n"
                    response_text += f"{'='*60}\n\n"
                    for i, notif in enumerate(notifications[:10], 1):  # Show max 10 announcements
                        notif_type = notif.get('type', 'info')
                        type_emoji = {'success': '✅', 'warning': '⚠️', 'error': '❌', 'info': 'ℹ️'}.get(notif_type, 'ℹ️')
                        is_read = '✓ Read' if notif.get('is_read', False) else '🆕 NEW'
                        
                        title = notif.get('title', 'Untitled')
                        msg = notif.get('message', '')
                        if len(msg) > 150:
                            msg = msg[:150] + '...'
                        date = notif.get('created_at', 'N/A')
                        
                        response_text += f"{type_emoji} **{title}** [{is_read}]\n"
                        response_text += f"   └─ 📝 Message:  {msg}\n"
                        response_text += f"   └─ 🕐 Date:     {date}\n"
                        if i < len(notifications[:10]):
                            response_text += "\n"
                    
                    response_text += f"\n{'='*60}"
                    data = {'notifications': notifications}
                else:
                    response_text = "✅ **Good news!** There are no announcements."
                    data = {}
            else:
                response_text = "Unable to fetch announcements at the moment."
                data = {}
        except Exception as e:
            response_text = "Unable to fetch announcements at the moment."
            data = {}
    elif intent in ['ask_admission', 'ask_fees', 'ask_cutoff', 'ask_college_info', 'ask_hostel', 'ask_transport', 'ask_placement', 'ask_scholarship', 'ask_documents']:
        # Handle college information queries using provided data
        response_text = get_college_info_response(msg, intent)
        data = {}
    elif intent == 'ask_guidance':
        # Delegate to the newly created personalized guidance function
        response_text = get_personalized_guidance(user_id, msg)
        data = {}
    else:
        response_text = "I can help with fees, attendance, marks, events, admissions, college information, and more. What would you like to know?"
        data = {}

    return jsonify({'intent': intent or 'unknown', 'response': response_text, 'data': data})

def get_personalized_guidance(user_id, msg):
    """Generate dynamic study suggestions using Groq API and student's marks."""
    if not user_id:
        return "Please log in to receive personalized academic guidance."
        
    student = Student.query.filter_by(user_id=user_id).first()
    if not student:
        return "You must be logged in as a student to get personalized guidance."

    # 1. Determine the subject to focus on
    target_subject = None
    
    # Check if a specific subject is mentioned in the message
    msg_lower = msg.lower()
    common_subjects = ['java', 'python', 'c++', 'cpp', 'dbms', 'sql', 'dsa', 'data structures', 'algorithms', 'math', 'machine learning', 'ai', 'operating systems', 'os', 'networks']
    for sub in common_subjects:
        if sub in msg_lower:
            target_subject = sub.title()
            break
            
    # If not explicitly stated, find their lowest mark
    if not target_subject:
        marks = Mark.query.filter_by(student_id=student.student_id).all()
        if marks:
            lowest_mark = min(marks, key=lambda x: x.obtained_marks)
            subject_obj = Subject.query.filter_by(subject_id=lowest_mark.subject_id).first()
            if subject_obj:
                target_subject = subject_obj.subject_name
                
    if not target_subject:
        target_subject = "your core subjects"

    # 2. Query Groq API
    from nlp_engine import nlp_engine
    if not nlp_engine.llm_configured or not getattr(nlp_engine, 'groq_client', None):
        return f"💡 **Study Tip for {target_subject}:** Practice with previous year question papers and break your study sessions into 25-minute Pomodoro intervals."
        
    prompt = (
        f"You are SmartEdu AI, a senior academic advisor for SKN Sinhgad College of Engineering. "
        f"A student is seeking specific guidance on {target_subject}. "
        "Provide a high-quality, personalized study plan including:\n"
        "1. **Core Concept Focus**: Identify 2-3 critical areas to master.\n"
        "2. **Learning Strategy**: Recommend a specific technique (e.g., active recall, spaced repetition).\n"
        "3. **Curated Resources**: List 2 specific, free high-quality resources (like NPTEL, Coursera, or top YouTube playlists).\n"
        "4. **Actionable Step**: One small task they can do *today*.\n\n"
        "Format in clean Markdown with professional headings. Keep the tone encouraging but rigorous. Max 180 words."
    )
    
    try:
        response = nlp_engine.groq_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.6,
            max_tokens=500
        )
        return f"✨ **SmartEdu personalized Roadmap: {target_subject}**\n\n" + response.choices[0].message.content
    except Exception as e:
        print(f"Groq guidance error: {e}")
        return f"💡 **Study Tip for {target_subject}:** Focus on understanding the core concepts and practice building small projects to reinforce your learning."

def get_marks_data(user_id, user_role, user):
    """Get marks data based on user role - admin sees all classes, student sees only their marks"""
    try:
        if user_role == 'admin':
            # Admin sees all classes marks
            marks_response = get_class_marks()
            data = json.loads(marks_response.get_data())
            return {'table': format_marks_table_for_admin(data)}
        else:
            # Student sees only their marks
            student = Student.query.filter_by(user_id=user_id).first()
            if student:
                marks = Mark.query.filter_by(student_id=student.student_id).all()
                return {'table': format_marks_table_for_student(student, marks)}
    except Exception as e:
        print(f"Error getting marks data: {str(e)}")
    return None

def get_attendance_data(user_id, user_role, user):
    """Get attendance data based on user role - admin sees all, student sees only their attendance"""
    try:
        if user_role == 'admin':
            # Admin sees all classes attendance
            attendance_response = get_class_attendance()
            data = json.loads(attendance_response.get_data())
            return {'table': format_attendance_table_for_admin(data)}
        else:
            # Student sees only their attendance
            student = Student.query.filter_by(user_id=user_id).first()
            if student:
                attendance = Attendance.query.filter_by(student_id=student.student_id).all()
                return {'table': format_attendance_table_for_student(student, attendance)}
    except Exception as e:
        print(f"Error getting attendance data: {str(e)}")
    return None

def get_fees_data(user_id, user_role, user):
    """Get fees data based on user role - admin sees all, student sees only their fees"""
    try:
        if user_role == 'admin':
            # Admin sees all students fees
            fees_response = get_class_fees()
            data = json.loads(fees_response.get_data())
            return {'table': format_fees_table_for_admin(data)}
        else:
            # Student sees only their fees
            student = Student.query.filter_by(user_id=user_id).first()
            if student and student.fees:
                return {'table': format_fees_table_for_student(student)}
    except Exception as e:
        print(f"Error getting fees data: {str(e)}")
    return None

def get_students_data(user_role):
    """Get students data - only admin can see all students"""
    if user_role == 'admin':
        try:
            students = Student.query.join(User).all()
            return {'table': format_students_table(students)}
        except Exception as e:
            print(f"Error getting students data: {str(e)}")
    return None

def format_marks_table_for_admin(class_data):
    """Format marks data for admin showing all classes"""
    if not class_data or len(class_data) == 0:
        return None
    table = {'title': 'Marks Overview - All Classes', 'columns': [
        {'key': 'class_name', 'label': 'Class'},
        {'key': 'roll_no', 'label': 'Roll No'},
        {'key': 'name', 'label': 'Name'},
        {'key': 'marks', 'label': 'Marks'},
        {'key': 'percentage', 'label': '%'}
    ], 'rows': []}
    for cls in class_data:
        for student in cls.get('students', []):
            table['rows'].append({
                'class_name': cls.get('class_name', ''),
                'roll_no': student.get('roll_number', ''),
                'name': student.get('name', ''),
                'marks': f"{student.get('total', 0)}/{student.get('total', 0)}",
                'percentage': f"{student.get('percentage', 0)}%"
            })
    return table

def format_marks_table_for_student(student, marks):
    """Format marks data for student showing only their marks"""
    if not marks:
        return None
    subjects = Subject.query.filter_by(class_id=student.class_id).all()
    table = {'title': 'Your Marks Overview', 'columns': [
        {'key': 'subject', 'label': 'Subject'},
        {'key': 'marks', 'label': 'Marks Obtained'},
        {'key': 'percentage', 'label': 'Percentage'}
    ], 'rows': []}
    for mark in marks:
        subject = Subject.query.get(mark.subject_id)
        if subject:
            table['rows'].append({
                'subject': subject.subject_name,
                'marks': f"{mark.obtained_marks}/35",
                'percentage': f"{(mark.obtained_marks / 35) * 100:.1f}%"
            })
    return table

def format_attendance_table_for_admin(class_data):
    """Format attendance data for admin showing all classes"""
    if not class_data or len(class_data) == 0:
        return None
    table = {'title': 'Attendance Overview - All Classes', 'columns': [
        {'key': 'class_name', 'label': 'Class'},
        {'key': 'roll_no', 'label': 'Roll No'},
        {'key': 'name', 'label': 'Name'},
        {'key': 'attendance', 'label': 'Attendance'},
        {'key': 'percentage', 'label': '%'}
    ], 'rows': []}
    for cls in class_data:
        for student in cls.get('students', []):
            table['rows'].append({
                'class_name': cls.get('class_name', ''),
                'roll_no': student.get('roll_number', ''),
                'name': student.get('name', ''),
                'attendance': f"{student.get('total', 0)}",
                'percentage': f"{student.get('total_percentage', 0)}%"
            })
    return table

def format_attendance_table_for_student(student, attendance):
    """Format attendance data for student showing only their attendance"""
    if not attendance:
        return None
    table = {'title': 'Your Attendance Overview', 'columns': [
        {'key': 'subject', 'label': 'Subject'},
        {'key': 'present', 'label': 'Present'},
        {'key': 'percentage', 'label': 'Attendance %'}
    ], 'rows': []}
    for att in attendance:
        subject = Subject.query.get(att.subject_id)
        if subject:
            percentage = (att.present_count / 50) * 100 if att.present_count else 0
            table['rows'].append({
                'subject': subject.subject_name,
                'present': f"{att.present_count}/50",
                'percentage': f"{percentage:.1f}%"
            })
    return table

def format_fees_table_for_admin(class_data):
    """Format fees data for admin showing all students"""
    if not class_data or len(class_data) == 0:
        return None
    table = {'title': 'Fees Overview - All Students', 'columns': [
        {'key': 'class_name', 'label': 'Class'},
        {'key': 'roll_no', 'label': 'Roll No'},
        {'key': 'name', 'label': 'Name'},
        {'key': 'total', 'label': 'Total Fees'},
        {'key': 'paid', 'label': 'Paid'},
        {'key': 'remaining', 'label': 'Remaining'},
        {'key': 'status', 'label': 'Status'}
    ], 'rows': []}
    for cls in class_data:
        for student in cls.get('students', []):
            table['rows'].append({
                'class_name': cls.get('class_name', ''),
                'roll_no': student.get('roll_number', ''),
                'name': student.get('name', ''),
                'total': f"₹{student.get('total_fees', 0)}",
                'paid': f"₹{student.get('paid_fees', 0)}",
                'remaining': f"₹{student.get('remaining_fees', 0)}",
                'status': student.get('payment_status', 'Unpaid')
            })
    return table

def format_fees_table_for_student(student):
    """Format fees data for student showing only their fees"""
    if not student.fees:
        return None
    table = {'title': 'Your Fee Details', 'columns': [
        {'key': 'fee_type', 'label': 'Fee Type'},
        {'key': 'total', 'label': 'Total Amount'},
        {'key': 'paid', 'label': 'Paid Amount'},
        {'key': 'remaining', 'label': 'Remaining'},
        {'key': 'status', 'label': 'Status'}
    ], 'rows': [{
        'fee_type': 'Total Fees',
        'total': f"₹{student.fees.total_amount}",
        'paid': f"₹{student.fees.paid_amount}",
        'remaining': f"₹{student.fees.due_amount}",
        'status': student.fees.payment_status
    }]}
    return table

def format_students_table(students):
    """Format students list for admin"""
    if not students:
        return None
    table = {'title': 'All Students', 'columns': [
        {'key': 'roll_no', 'label': 'Roll No'},
        {'key': 'name', 'label': 'Name'},
        {'key': 'email', 'label': 'Email'},
        {'key': 'contact', 'label': 'Contact'},
        {'key': 'class', 'label': 'Class'}
    ], 'rows': []}
    for student in students:
        table['rows'].append({
            'roll_no': student.roll_no,
            'name': student.user.name,
            'email': student.user.email,
            'contact': student.user.contact_no or '-',
            'class': student.class_.class_name
        })
    return table

# Routes
@app.route('/')
def home():
    return jsonify({'message': 'SmartEdu Chatbot API is running!'})

@app.route('/api/users', methods=['GET'])
def get_users():
    users = User.query.all()
    return jsonify([user.to_dict() for user in users])

@app.route('/api/users', methods=['POST'])
def create_user():
    data = request.get_json()
    
    if not data or not data.get('username') or not data.get('email'):
        return jsonify({'error': 'Username and email are required'}), 400
    
    user = User(username=data['username'], email=data['email'])
    db.session.add(user)
    db.session.commit()
    
    return jsonify(user.to_dict()), 201

@app.route('/chat', methods=['POST'])
def chat():
    """
    Process user message with Dialogflow intent detection and sentiment analysis
    Uses the detect_intent_texts function as requested
    Enhanced to support role-based data access
    """
    try:
        data = request.get_json()
        
        if not data or not data.get('message'):
            return jsonify({'error': 'Message is required'}), 400
        
        user_message = data['message']
        user_id = data.get('user_id', 1)  # Default user if not provided
        language_code = data.get('language_code', 'en')  # Default to English
        
        # Get user details to determine role
        # Use SQLAlchemy 2.0 style Session.get
        user = db.session.get(User, user_id)
        user_role = user.role if user else 'student'
        
        # Analyze sentiment of user message
        sentiment_data = analyze_sentiment(user_message)
        
        # Get project_id and session_id from request or use defaults
        project_id = data.get('project_id', dialogflow_service.project_id)
        session_id = data.get('session_id', f'smartedu-session-{user_id}')
        
        # Use Rasa NLU for intent detection
        dialogflow_result = rasa_service.detect_intent_texts(
            text=user_message,
            language_code=language_code
        )
        
        # Get educational response based on intent
        # Build educational response via local helper (no Dialogflow dependency)
        educational_response = get_educational_response(
            dialogflow_result['intent'],
            dialogflow_result['parameters'],
            user_message,
            dialogflow_result['fulfillment_text']
        )
        
        # Enhance response with role-based data based on intent
        data_response = None
        if 'marks' in dialogflow_result['intent'].lower():
            data_response = get_marks_data(user_id, user_role, user)
        elif 'attendance' in dialogflow_result['intent'].lower():
            data_response = get_attendance_data(user_id, user_role, user)
        elif 'fees' in dialogflow_result['intent'].lower():
            data_response = get_fees_data(user_id, user_role, user)
        elif 'student' in dialogflow_result['intent'].lower():
            data_response = get_students_data(user_role)
        
        # Add empathetic prefix if sentiment is negative
        empathetic_prefix = get_empathetic_prefix(sentiment_data)
        final_response = empathetic_prefix + educational_response['response']
        
        # Create chat message record
        chat_message = ChatMessage(
            user_id=user_id,
            message=user_message,
            response=final_response,
            intent=dialogflow_result['intent'],
            confidence=dialogflow_result['confidence']
        )
        
        db.session.add(chat_message)
        db.session.commit()
        
        # Prepare comprehensive response for frontend
        response_data = {
            'message_id': chat_message.id,
            'response': final_response,
            'intent': dialogflow_result['intent'],
            'confidence': dialogflow_result['confidence'],
            'suggestions': educational_response.get('suggestions', []),
            'parameters': dialogflow_result.get('parameters', {}),
            'action': dialogflow_result.get('action', ''),
            'all_required_params_present': dialogflow_result.get('all_required_params_present', False),
            'query_text': dialogflow_result.get('query_text', user_message),
            'timestamp': chat_message.timestamp.isoformat(),
            'session_id': session_id,
            'project_id': project_id,
            'sentiment': {
                'polarity': sentiment_data['polarity'],
                'sentiment': sentiment_data['sentiment'],
                'has_empathetic_prefix': bool(empathetic_prefix)
            }
        }
        
        # Add table data if available
        if data_response and data_response.get('table'):
            response_data['table'] = data_response['table']
        
        return jsonify(response_data), 200
        
    except Exception as e:
        return jsonify({'error': f'Chat processing failed: {str(e)}'}), 500

@app.route('/events', methods=['GET'])
def get_events():
    """
    Fetch event list from database
    """
    try:
        # Get query parameters
        active_only = request.args.get('active_only', 'true').lower() == 'true'
        event_type = request.args.get('type')
        limit = request.args.get('limit', type=int)
        
        # Build query
        query = Event.query
        
        if active_only:
            query = query.filter(Event.is_active == True)
            # Also filter out past events
            today = datetime.utcnow().date()
            query = query.filter(Event.event_date >= today)
        
        if event_type:
            query = query.filter(Event.event_type == event_type)
        
        # Order by event date
        query = query.order_by(Event.event_date.asc())
        
        if limit:
            query = query.limit(limit)
        
        events = query.all()
        
        return jsonify([event.to_dict() for event in events]), 200
        
    except Exception as e:
        return jsonify({'error': f'Failed to fetch events: {str(e)}'}), 500

@app.route('/admin/announce', methods=['POST'])
def create_announcement():
    """
    Add new announcement (admin only)
    """
    try:
        data = request.get_json()
        
        if not data or not data.get('title') or not data.get('message'):
            return jsonify({'error': 'Title and message are required'}), 400
        
        # Parse optional fields
        priority = data.get('priority', 'normal')
        expires_at = data.get('expires_at')
        
        if expires_at:
            try:
                expires_at = datetime.fromisoformat(expires_at.replace('Z', '+00:00'))
            except ValueError:
                return jsonify({'error': 'Invalid expires_at format. Use ISO format.'}), 400
        
        # Create announcement
        announcement = Announcement(
            title=data['title'],
            message=data['message'],
            priority=priority,
            expires_at=expires_at
        )
        
        db.session.add(announcement)
        db.session.commit()
        
        return jsonify(announcement.to_dict()), 201
        
    except Exception as e:
        return jsonify({'error': f'Failed to create announcement: {str(e)}'}), 500

@app.route('/api/announcements', methods=['GET', 'POST'])
def get_announcements():
    """
    Get active announcements or create new announcement
    """
    if request.method == 'POST':
        try:
            data = request.get_json()
            
            if not data or not data.get('title') or not data.get('body'):
                return jsonify({'error': 'Title and body are required'}), 400
            
            # Create announcement
            announcement = Announcement(
                title=data.get('title'),
                message=data.get('body'),  # Use 'body' from frontend
                priority=data.get('type', 'normal'),  # Use 'type' from frontend
                target=data.get('target_audience', 'all'),
                is_active=True,
                created_at=datetime.utcnow()
            )
            
            db.session.add(announcement)
            db.session.commit()
            
            return jsonify(announcement.to_dict()), 201
            
        except Exception as e:
            db.session.rollback()
            return jsonify({'error': f'Failed to create announcement: {str(e)}'}), 500
    else:
        # GET request
        try:
            active_only = request.args.get('active_only', 'true').lower() == 'true'
            priority = request.args.get('priority')
            
            query = Announcement.query
            
            if active_only:
                query = query.filter(Announcement.is_active == True)
                # Filter out expired announcements
                query = query.filter(
                    (Announcement.expires_at.is_(None)) | 
                    (Announcement.expires_at > datetime.utcnow())
                )
            
            if priority:
                query = query.filter(Announcement.priority == priority)
            
            query = query.order_by(Announcement.created_at.desc())
            
            announcements = query.all()
            
            return jsonify([announcement.to_dict() for announcement in announcements]), 200
            
        except Exception as e:
            return jsonify({'error': f'Failed to fetch announcements: {str(e)}'}), 500

@app.route('/api/chat/<int:user_id>', methods=['GET'])
def get_user_chats(user_id):
    chats = ChatMessage.query.filter_by(user_id=user_id).order_by(ChatMessage.timestamp.desc()).all()
    return jsonify([chat.to_dict() for chat in chats])

# ==================== NEW API ENDPOINTS ====================

# DEPARTMENTS
@app.route('/api/departments', methods=['GET'])
def get_departments():
    """Get all departments"""
    try:
        departments = Department.query.filter_by(is_active=True).all()
        return jsonify([dept.to_dict() for dept in departments]), 200
    except Exception as e:
        return jsonify({'error': f'Failed to fetch departments: {str(e)}'}), 500

# CLASSES
@app.route('/api/classes', methods=['GET'])
def get_classes():
    """Get all classes with optional department filter"""
    try:
        dept_id = request.args.get('dept_id', type=int)
        query = Class.query.filter_by(is_active=True)
        
        if dept_id:
            query = query.filter_by(dept_id=dept_id)
        
        classes = query.all()
        return jsonify([cls.to_dict() for cls in classes]), 200
    except Exception as e:
        return jsonify({'error': f'Failed to fetch classes: {str(e)}'}), 500

# SUBJECTS
@app.route('/api/subjects', methods=['GET'])
def get_subjects():
    """Get all subjects with optional class filter"""
    try:
        class_id = request.args.get('class_id', type=int)
        query = Subject.query.filter_by(is_active=True)
        
        if class_id:
            query = query.filter_by(class_id=class_id)
        
        subjects = query.all()
        return jsonify([subject.to_dict() for subject in subjects]), 200
    except Exception as e:
        return jsonify({'error': f'Failed to fetch subjects: {str(e)}'}), 500

# STUDENTS
@app.route('/api/students', methods=['GET'])
def get_students():
    """Get all students with optional filters (admin only)"""
    try:
        # Only allow non-student roles to list all students
        user, error = get_current_user_from_token()
        if error:
            return error
        if user.role == 'student':
            return jsonify({'error': 'Not authorized to view all students'}), 403

        class_id = request.args.get('class_id', type=int)
        roll_no = request.args.get('roll_no')
        limit = request.args.get('limit', type=int)
        query = Student.query.filter_by(is_active=True)
        
        if class_id:
            query = query.filter_by(class_id=class_id)
        if roll_no:
            query = query.filter_by(roll_no=roll_no)

        # Apply limit if provided to avoid returning very large result sets
        if limit and limit > 0:
            query = query.limit(limit)
        
        students = query.all()
        return jsonify([student.to_dict() for student in students]), 200
    except Exception as e:
        return jsonify({'error': f'Failed to fetch students: {str(e)}'}), 500


@app.route('/api/students/me', methods=['GET'])
def get_my_student_record():
    """
    Get the current logged-in student's record using the users->students mapping
    defined in schema.txt (students.user_id is UNIQUE and references users.user_id).
    """
    try:
        user, error = get_current_user_from_token()
        if error:
            return error

        if user.role != 'student':
            return jsonify({'error': 'Only students have a personal student record'}), 403

        student = Student.query.filter_by(user_id=user.user_id, is_active=True).first()
        if not student:
            return jsonify({'error': 'Student record not found for current user'}), 404

        return jsonify(student.to_dict()), 200
    except Exception as e:
        return jsonify({'error': f'Failed to fetch current student: {str(e)}'}), 500

@app.route('/api/students/me/dashboard-stats', methods=['GET'])
def get_my_dashboard_stats():
    """
    Get aggregated progress statistics (Marks, Attendance, Fees) for the logged-in student's dashboard.
    """
    try:
        user, error = get_current_user_from_token()
        if error:
            return error

        if user.role != 'student':
            return jsonify({'error': 'Only students can view personal progress stats'}), 403

        student = Student.query.filter_by(user_id=user.user_id, is_active=True).first()
        if not student:
            return jsonify({'error': 'Student record not found'}), 404

        # 1. Marks Percentage
        marks = Mark.query.filter_by(student_id=student.student_id).all()
        total_obtained = sum(m.obtained_marks or 0 for m in marks)
        total_max = sum(m.total_marks or 35 for m in marks)
        marks_pct = (total_obtained / total_max * 100) if total_max > 0 else 0

        # 2. Attendance Percentage
        attendances = Attendance.query.filter_by(student_id=student.student_id).all()
        total_present = sum(a.present_count or 0 for a in attendances)
        total_classes = sum(a.total_classes or 0 for a in attendances)
        attendance_pct = (total_present / total_classes * 100) if total_classes > 0 else 0

        # 3. Fees Percentage
        fee_record = Fee.query.filter_by(student_id=student.student_id).first()
        fees_pct = (float(fee_record.paid_amount) / float(fee_record.total_amount) * 100) if fee_record and fee_record.total_amount and float(fee_record.total_amount) > 0 else 0
        if fee_record and fee_record.payment_status == 'Paid':
             fees_pct = 100

        return jsonify({
            'marks_percentage': round(marks_pct, 1),
            'attendance_percentage': round(attendance_pct, 1),
            'fees_percentage': round(fees_pct, 1),
        }), 200

    except Exception as e:
        return jsonify({'error': f'Failed to fetch dashboard stats: {str(e)}'}), 500

@app.route('/api/students', methods=['POST'])
def create_student():
    """Create a new student"""
    try:
        data = request.get_json()
        
        required_fields = ['user_id', 'roll_no', 'class_id', 'admission_year']
        for field in required_fields:
            if not data.get(field):
                return jsonify({'error': f'{field} is required'}), 400
        
        # Check if user exists and is a student
        user = db.session.get(User, data['user_id'])
        if not user or user.role != 'student':
            return jsonify({'error': 'Invalid user or user is not a student'}), 400
        
        # Check if roll number already exists in this class only (same roll allowed in different classes)
        if Student.query.filter_by(roll_no=data['roll_no'], class_id=data['class_id']).first():
            return jsonify({'error': f'Roll number {data["roll_no"]} already exists in this class'}), 400
        
        student = Student(
            user_id=data['user_id'],
            roll_no=data['roll_no'],
            class_id=data['class_id'],
            admission_year=data['admission_year'],
            admission_date=datetime.strptime(data.get('admission_date', datetime.now().strftime('%Y-%m-%d')), '%Y-%m-%d').date() if data.get('admission_date') else None,
            guardian_name=data.get('guardian_name'),
            guardian_contact=data.get('guardian_contact'),
            address=data.get('address')
        )
        
        db.session.add(student)
        db.session.flush()  # get student_id

        # Ensure the class has subjects so we can seed marks/attendance with 0 values
        cls = db.session.get(Class, student.class_id)
        subjects = Subject.query.filter_by(class_id=student.class_id).all()
        if not subjects and cls:
            for i in range(1, 6):
                db.session.add(Subject(
                    class_id=student.class_id,
                    subject_name=f"Subject {i}",
                    subject_code=f"{cls.class_code}-S{i}"[:20],
                    credits=4,
                    description="",
                    is_active=True,
                    created_at=datetime.utcnow()
                ))
            db.session.flush()
            subjects = Subject.query.filter_by(class_id=student.class_id).all()

        # Seed default marks/attendance (0) for all subjects
        for subject in subjects:
            if not Mark.query.filter_by(student_id=student.student_id, subject_id=subject.subject_id).first():
                db.session.add(Mark(
                    student_id=student.student_id,
                    subject_id=subject.subject_id,
                    total_marks=35,
                    obtained_marks=0,
                    exam_date=datetime.now().date(),
                    created_at=datetime.utcnow(),
                    updated_at=datetime.utcnow()
                ))
            if not Attendance.query.filter_by(student_id=student.student_id, subject_id=subject.subject_id).first():
                db.session.add(Attendance(
                    student_id=student.student_id,
                    subject_id=subject.subject_id,
                    present_count=0,
                    absent_count=0,
                    late_count=0,
                    total_classes=50,
                    attendance_percentage=0.0,
                    academic_year=str(datetime.now().year),
                    updated_at=datetime.utcnow()
                ))

        # Seed default fees row if missing
        if not Fee.query.filter_by(student_id=student.student_id).first():
            db.session.add(Fee(
                student_id=student.student_id,
                total_amount=50000.00,
                paid_amount=0.00,
                due_amount=50000.00,
                payment_status='Unpaid',
                created_at=datetime.utcnow(),
                updated_at=datetime.utcnow()
            ))

        db.session.commit()
        
        return jsonify(student.to_dict()), 201
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'Failed to create student: {str(e)}'}), 500

@app.route('/api/students/<int:student_id>', methods=['GET'])
def get_student(student_id):
    """Get student by ID with detailed information"""
    try:
        student = Student.query.get_or_404(student_id)
        student_data = student.to_dict()
        
        # Add user information
        student_data['user'] = student.user.to_dict()
        
        # Add class information
        student_data['class'] = student.class_.to_dict()
        
        # Add fees information
        if student.fees:
            student_data['fees'] = student.fees.to_dict()
        
        return jsonify(student_data), 200
    except Exception as e:
        return jsonify({'error': f'Failed to fetch student: {str(e)}'}), 500

# MARKS
@app.route('/api/students/<int:student_id>/marks', methods=['GET'])
def get_student_marks(student_id):
    """Get marks for a specific student with subject names"""
    try:
        student = Student.query.get_or_404(student_id)
        subject_id = request.args.get('subject_id', type=int)
        
        query = Mark.query.filter_by(student_id=student_id)
        
        if subject_id:
            query = query.filter_by(subject_id=subject_id)
        
        marks = query.all()
        
        # Enhance marks data with subject information
        marks_data = []
        for mark in marks:
            mark_dict = mark.to_dict()
            # Add subject information
            subject = db.session.get(Subject, mark.subject_id)
            if subject:
                mark_dict['subject_name'] = subject.subject_name
                mark_dict['subject_code'] = subject.subject_code
                mark_dict['credits'] = subject.credits
            marks_data.append(mark_dict)
        
        return jsonify(marks_data), 200
    except Exception as e:
        return jsonify({'error': f'Failed to fetch marks: {str(e)}'}), 500

@app.route('/api/students/<int:student_id>/marks', methods=['POST'])
def add_student_marks(student_id):
    """Add marks for a student"""
    try:
        data = request.get_json()
        
        required_fields = ['subject_id', 'obtained_marks']
        for field in required_fields:
            if not data.get(field):
                return jsonify({'error': f'{field} is required'}), 400
        
        # Check if student exists
        student = Student.query.get_or_404(student_id)
        
        # Check if subject exists
        subject = Subject.query.get_or_404(data['subject_id'])
        
        # Check if marks already exist for this student-subject combination
        existing_mark = Mark.query.filter_by(
            student_id=student_id,
            subject_id=data['subject_id']
        ).first()
        
        if existing_mark:
            return jsonify({'error': 'Marks already exist for this subject'}), 400
        
        obtained_marks = data['obtained_marks']
        total_marks = data.get('total_marks', 35)
        
        # Validate marks range
        if obtained_marks < 0 or obtained_marks > 35:
            return jsonify({'error': 'Marks must be between 0 and 35.'}), 400
        
        # Validate maximum marks
        if obtained_marks > total_marks:
            return jsonify({'error': f'Marks cannot exceed total marks ({total_marks})'}), 400
        
        # Create mark record
        mark = Mark(
            student_id=student_id,
            subject_id=data['subject_id'],
            total_marks=total_marks,
            obtained_marks=obtained_marks,
            exam_date=datetime.strptime(data.get('exam_date', datetime.now().strftime('%Y-%m-%d')), '%Y-%m-%d').date() if data.get('exam_date') else None
        )
        
        db.session.add(mark)
        db.session.commit()
        
        return jsonify(mark.to_dict()), 201
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'Failed to add marks: {str(e)}'}), 500

@app.route('/api/students/<int:student_id>/marks/<int:mark_id>', methods=['PUT'])
def update_student_marks(student_id, mark_id):
    """Update marks for a student"""
    try:
        data = request.get_json()
        
        # Check if mark exists
        mark = Mark.query.filter_by(mark_id=mark_id, student_id=student_id).first()
        if not mark:
            return jsonify({'error': 'Mark not found'}), 404
        
        # Validate minimum marks requirement
        if 'obtained_marks' in data:
            obtained_marks = data['obtained_marks']
            total_marks = data.get('total_marks', mark.total_marks)
            
            if obtained_marks < 0 or obtained_marks > 35:
                return jsonify({'error': 'Marks must be between 0 and 35.'}), 400
            
            if obtained_marks > total_marks:
                return jsonify({'error': f'Marks cannot exceed total marks ({total_marks})'}), 400
            
            # Update marks
            mark.obtained_marks = obtained_marks
            mark.total_marks = total_marks
        
        # Update other fields if provided
        if 'exam_date' in data:
            mark.exam_date = datetime.strptime(data['exam_date'], '%Y-%m-%d').date()
        
        mark.updated_at = datetime.utcnow()
        
        db.session.commit()
        return jsonify(mark.to_dict()), 200
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'Failed to update marks: {str(e)}'}), 500


# ATTENDANCE
@app.route('/api/students/<int:student_id>/attendance', methods=['GET'])
def get_student_attendance(student_id):
    """Get attendance for a specific student"""
    try:
        student = Student.query.get_or_404(student_id)
        summary = request.args.get('summary', 'false').lower() == 'true'
        subject_id = request.args.get('subject_id', type=int)
        # Note: Attendance model stores aggregated counts (present_count, total_classes, etc.)
        # so detailed date-range filtering isn't supported by the current schema.
        # Ignore start_date/end_date to avoid referencing non-existent columns.
        start_date = request.args.get('start_date')
        end_date = request.args.get('end_date')
        
        if summary:
            # Return attendance summary grouped by subject
            query = Attendance.query.filter_by(student_id=student_id)
            
            if start_date:
                query = query.filter(Attendance.date >= datetime.strptime(start_date, '%Y-%m-%d').date())
            if end_date:
                query = query.filter(Attendance.date <= datetime.strptime(end_date, '%Y-%m-%d').date())
            
            attendance_records = query.all()
            
            # Group by subject and calculate summary
            subject_summary = {}
            for att in attendance_records:
                if att.subject_id not in subject_summary:
                    subject_summary[att.subject_id] = {
                        'subject_id': att.subject_id,
                        'total_classes': 0,
                        'present_count': 0,
                        'absent_count': 0
                    }
                
                subject_summary[att.subject_id]['total_classes'] += 1
                if att.status == 'Present':
                    subject_summary[att.subject_id]['present_count'] += 1
                else:
                    subject_summary[att.subject_id]['absent_count'] += 1
            
            # Fetch subject details and format response
            summary_list = []
            for subj_id, summary_data in subject_summary.items():
                subject = db.session.get(Subject, subj_id)
                if subject:
                    summary_list.append({
                        'subject_id': subj_id,
                        'subject_name': subject.subject_name,
                        'subject_code': subject.subject_code,
                        'present_count': summary_data['present_count'],
                        'absent_count': summary_data['absent_count'],
                        'total_classes': summary_data['total_classes'],
                        'attendance_percentage': round((summary_data['present_count'] / summary_data['total_classes']) * 100, 2) if summary_data['total_classes'] > 0 else 0
                    })
            
            return jsonify(summary_list), 200
        else:
            # Return aggregated attendance records for the student (no per-date records)
            query = Attendance.query.filter_by(student_id=student_id)

            if subject_id:
                query = query.filter_by(subject_id=subject_id)

            attendance = query.order_by(Attendance.updated_at.desc()).all()
            return jsonify([att.to_dict() for att in attendance]), 200
    except Exception as e:
        return jsonify({'error': f'Failed to fetch attendance: {str(e)}'}), 500

@app.route('/api/students/<int:student_id>/attendance', methods=['POST'])
def mark_attendance(student_id):
    """Mark attendance for a student"""
    try:
        data = request.get_json()
        
        required_fields = ['subject_id', 'status']
        for field in required_fields:
            if not data.get(field):
                return jsonify({'error': f'{field} is required'}), 400
        
        # Check if student exists
        student = Student.query.get_or_404(student_id)
        
        # Check if subject exists
        subject = Subject.query.get_or_404(data['subject_id'])
        
        attendance_date = datetime.strptime(data.get('date', datetime.now().strftime('%Y-%m-%d')), '%Y-%m-%d').date()
        
        # Check if attendance already marked for this date
        existing_attendance = Attendance.query.filter_by(
            student_id=student_id,
            subject_id=data['subject_id'],
            date=attendance_date
        ).first()
        
        if existing_attendance:
            # Update existing attendance
            existing_attendance.status = data['status']
            existing_attendance.remarks = data.get('remarks')
            db.session.commit()
            return jsonify(existing_attendance.to_dict()), 200
        else:
            # Create new attendance record
            attendance = Attendance(
                student_id=student_id,
                subject_id=data['subject_id'],
                date=attendance_date,
                status=data['status'],
                remarks=data.get('remarks')
            )
            
            db.session.add(attendance)
            db.session.commit()
            
            return jsonify(attendance.to_dict()), 201
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'Failed to mark attendance: {str(e)}'}), 500

# FEES
@app.route('/api/students/<int:student_id>/attendance-summary', methods=['GET'])
def get_student_attendance_summary(student_id):
    """Get attendance summary for a student (aggregated counts)"""
    try:
        student = Student.query.get_or_404(student_id)
        
        summaries = Attendance.query.filter_by(student_id=student_id).all()
        
        # Get subject details for each summary
        summary_list = []
        for summary in summaries:
            subject = db.session.get(Subject, summary.subject_id)
            if subject:
                summary_list.append({
                    'attendance_id': summary.attendance_id,
                    'subject_id': subject.subject_id,
                    'subject_name': subject.subject_name,
                    'subject_code': subject.subject_code,
                    'present_count': summary.present_count,
                    'absent_count': summary.absent_count,
                    'late_count': summary.late_count,
                    'total_classes': summary.total_classes,
                    'attendance_percentage': float(summary.attendance_percentage),
                    'academic_year': summary.academic_year
                })
        
        return jsonify(summary_list), 200
    except Exception as e:
        return jsonify({'error': f'Failed to fetch attendance summary: {str(e)}'}), 500

@app.route('/api/students/<int:student_id>/fees', methods=['GET'])
def get_student_fees(student_id):
    """Get fees information for a student"""
    try:
        student = Student.query.get_or_404(student_id)
        
        if not student.fees:
            return jsonify({'error': 'No fees record found for this student'}), 404
        
        return jsonify(student.fees.to_dict()), 200
    except Exception as e:
        return jsonify({'error': f'Failed to fetch fees: {str(e)}'}), 500

@app.route('/api/students/<int:student_id>/fees', methods=['PUT'])
def update_student_fees(student_id):
    """Update fees payment for a student"""
    try:
        data = request.get_json()
        
        student = Student.query.get_or_404(student_id)
        
        if not student.fees:
            return jsonify({'error': 'No fees record found for this student'}), 404
        
        if 'paid_amount' in data:
            student.fees.paid_amount = data['paid_amount']
            student.fees.due_amount = student.fees.total_amount - student.fees.paid_amount
            student.fees.last_payment_date = date.today()
            
            # Update payment status
            if student.fees.paid_amount == 0:
                student.fees.payment_status = 'Unpaid'
            elif student.fees.paid_amount >= student.fees.total_amount:
                student.fees.payment_status = 'Paid'
            else:
                student.fees.payment_status = 'Partial'
        
        db.session.commit()
        return jsonify(student.fees.to_dict()), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'Failed to update fees: {str(e)}'}), 500

# NOTIFICATIONS
@app.route('/api/users/<int:user_id>/notifications', methods=['GET'])
def get_user_notifications(user_id):
    """Get notifications for a user"""
    try:
        user = User.query.get_or_404(user_id)
        unread_only = request.args.get('unread_only', 'false').lower() == 'true'
        
        query = Notification.query.filter_by(user_id=user_id)
        
        if unread_only:
            query = query.filter_by(is_read=False)
        
        notifications = query.order_by(Notification.created_at.desc()).all()
        return jsonify([notif.to_dict() for notif in notifications]), 200
    except Exception as e:
        return jsonify({'error': f'Failed to fetch notifications: {str(e)}'}), 500

@app.route('/api/notifications', methods=['GET'])
def get_notifications():
    """Get notifications for current user (requires authentication)"""
    try:
        # In production, verify JWT token here
        auth_header = request.headers.get('Authorization')
        if not auth_header or not auth_header.startswith('Bearer '):
            return jsonify({'error': 'Authorization token required'}), 401
        
        token = auth_header.split(' ')[1]
        # For now, extract user_id from token (in production, verify JWT)
        if token.startswith('token-'):
            user_id = int(token.split('-')[1])
            user = db.session.get(User, user_id)
            
            if not user:
                return jsonify({'error': 'User not found'}), 404
            
            unread_only = request.args.get('unread_only', 'false').lower() == 'true'
            
            query = Notification.query.filter_by(user_id=user_id)
            
            if unread_only:
                query = query.filter_by(is_read=False)
            
            notifications = query.order_by(Notification.created_at.desc()).all()
            return jsonify([notif.to_dict() for notif in notifications]), 200
        else:
            return jsonify({'error': 'Invalid token'}), 401
            
    except Exception as e:
        return jsonify({'error': f'Failed to fetch notifications: {str(e)}'}), 500

@app.route('/api/notifications/unread-count', methods=['GET'])
def get_unread_notifications_count():
    """Get unread notifications count for current user"""
    try:
        # In production, verify JWT token here
        auth_header = request.headers.get('Authorization')
        if not auth_header or not auth_header.startswith('Bearer '):
            return jsonify({'error': 'Authorization token required'}), 401
        
        token = auth_header.split(' ')[1]
        # For now, extract user_id from token (in production, verify JWT)
        if token.startswith('token-'):
            user_id = int(token.split('-')[1])
            user = db.session.get(User, user_id)
            
            if not user:
                return jsonify({'error': 'User not found'}), 404
            
            unread_count = Notification.query.filter_by(user_id=user_id, is_read=False).count()
            return jsonify({'unread_count': unread_count}), 200
        else:
            return jsonify({'error': 'Invalid token'}), 401
            
    except Exception as e:
        return jsonify({'error': f'Failed to fetch unread count: {str(e)}'}), 500

@app.route('/api/notifications/<int:notification_id>/read', methods=['PUT'])
def mark_notification_read(notification_id):
    """Mark a notification as read"""
    try:
        notification = Notification.query.get_or_404(notification_id)
        notification.is_read = True
        db.session.commit()
        
        return jsonify(notification.to_dict()), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'Failed to mark notification as read: {str(e)}'}), 500

# DASHBOARD STATS
@app.route('/api/dashboard/stats', methods=['GET'])
def get_dashboard_stats():
    """Get dashboard statistics"""
    try:
        stats = {
            'total_students': Student.query.filter_by(is_active=True).count(),
            'total_classes': Class.query.filter_by(is_active=True).count(),
            'total_subjects': Subject.query.filter_by(is_active=True).count(),
            'total_events': Event.query.filter_by(is_active=True).count(),
            'upcoming_events': Event.query.filter(
                Event.event_date >= date.today(),
                Event.is_active == True
            ).count(),
            'unpaid_fees': Fee.query.filter_by(payment_status='Unpaid').count(),
            'partial_fees': Fee.query.filter_by(payment_status='Partial').count()
        }
        
        return jsonify(stats), 200
    except Exception as e:
        return jsonify({'error': f'Failed to fetch dashboard stats: {str(e)}'}), 500

# STUDENT SERVICES - CLASS-BASED DATA ENDPOINTS

@app.route('/api/student-services/marks', methods=['GET'])
def get_class_marks():
    """Get marks data grouped by class for student services"""
    try:
        user, error = get_current_user_from_token()
        if error: return error

        if user.role == 'student':
            student = Student.query.filter_by(user_id=user.user_id, is_active=True).first()
            if not student: return jsonify({'error': 'Student record not found'}), 404
            students = [student]
        else:
            students = Student.query.join(Class).join(User).all()
        
        class_data = {}
        for student in students:
            class_name = student.class_.class_name
            class_id = student.class_id
            if class_name not in class_data:
                subjects = Subject.query.filter_by(class_id=class_id, is_active=True).all()
                class_data[class_name] = {
                    'class_name': class_name,
                    'class_id': class_id,
                    'subjects': [s.to_dict() for s in subjects],
                    'students': []
                }
            
            marks = Mark.query.filter_by(student_id=student.student_id).all()
            marks_map = {m.subject_id: {'obtained': m.obtained_marks, 'total': m.total_marks or 35} for m in marks}
            
            student_subject_marks = {}
            total_sum = 0
            max_sum = 0
            subjects = Subject.query.filter_by(class_id=class_id, is_active=True).all()
            for s in subjects:
                # Use Subject-level total_marks (official)
                s_total = s.total_marks or 35
                m_info = marks_map.get(s.subject_id, {'obtained': 0, 'total': s_total})
                # Overwrite cached total with official total just in case
                m_info['total'] = s_total
                student_subject_marks[str(s.subject_id)] = m_info
                total_sum += m_info['obtained']
                max_sum += s_total
                
            student_marks = {
                'student_id': student.student_id,
                'roll_number': student.roll_no,
                'name': student.user.name,
                'marks': student_subject_marks,
                'total': total_sum,
                'max_total': max_sum,
                'percentage': round((total_sum / max_sum * 100), 1) if max_sum > 0 else 0.0
            }
            
            # For backward compatibility
            for i, s in enumerate(subjects[:5]):
                n = i + 1
                s_total = s.total_marks or 35
                m_info = marks_map.get(s.subject_id, {'obtained': 0, 'total': s_total})
                student_marks[f'sub{n}'] = m_info['obtained']
                student_marks[f'sub{n}_total'] = s_total
                student_marks[f'sub{n}_name'] = s.subject_name
                student_marks[f'sub{n}_subject_id'] = s.subject_id
            
            class_data[class_name]['students'].append(student_marks)
        
        return jsonify(list(class_data.values())), 200
    except Exception as e:
        return jsonify({'error': f'Failed to fetch marks: {str(e)}'}), 500

@app.route('/api/student-services/attendance', methods=['GET'])
def get_class_attendance():
    """Get attendance data grouped by class for student services"""
    try:
        user, error = get_current_user_from_token()
        if error: return error

        if user.role == 'student':
            student = Student.query.filter_by(user_id=user.user_id, is_active=True).first()
            if not student: return jsonify({'error': 'Student record not found'}), 404
            students = [student]
        else:
            students = Student.query.join(Class).join(User).all()
        
        class_data = {}
        for student in students:
            class_name = student.class_.class_name
            class_id = student.class_id
            if class_name not in class_data:
                subjects = Subject.query.filter_by(class_id=class_id, is_active=True).all()
                class_data[class_name] = {
                    'class_name': class_name,
                    'class_id': class_id,
                    'subjects': [s.to_dict() for s in subjects],
                    'students': []
                }
            
            attendance_records = Attendance.query.filter_by(student_id=student.student_id).all()
            att_map = {a.subject_id: a for a in attendance_records}
            
            student_subject_attendance = {}
            total_present = 0
            total_classes_all = 0
            subjects = Subject.query.filter_by(class_id=class_id, is_active=True).all()
            
            for s in subjects:
                att = att_map.get(s.subject_id)
                present = att.present_count if att else 0
                # Use Subject-level total_classes
                s_total = s.total_classes or 50
                student_subject_attendance[str(s.subject_id)] = {
                    'present': present,
                    'total': s_total,
                    'percentage': float(att.attendance_percentage) if att and att.attendance_percentage else 0.0
                }
                total_present += present
                total_classes_all += s_total
            
            student_attendance = {
                'student_id': student.student_id,
                'roll_number': student.roll_no,
                'name': student.user.name,
                'attendance': student_subject_attendance,
                'total_present': total_present,
                'total_classes': total_classes_all,
                'total_percentage': round((total_present / total_classes_all * 100), 1) if total_classes_all > 0 else 0.0,
                'is_defaulter': (total_present / total_classes_all * 100) < 75.0 if total_classes_all > 0 else True
            }

            # Backward compatibility
            for i, s in enumerate(subjects[:5]):
                n = i + 1
                att = att_map.get(s.subject_id)
                s_total = s.total_classes or 50
                student_attendance[f'sub{n}'] = att.present_count if att else 0
                student_attendance[f'sub{n}_name'] = s.subject_name
                student_attendance[f'sub{n}_total'] = s_total

            class_data[class_name]['students'].append(student_attendance)
        
        return jsonify(list(class_data.values())), 200
    except Exception as e:
        return jsonify({'error': f'Failed to fetch attendance: {str(e)}'}), 500

@app.route('/api/student-services/fees', methods=['GET'])
def get_class_fees():
    """Get fees data grouped by class for student services"""
    try:
        # Determine current user and scope data appropriately
        user, error = get_current_user_from_token()
        if error:
            return error

        # For students, only return their own fees; for staff, return all
        if user.role == 'student':
            student = Student.query.filter_by(user_id=user.user_id, is_active=True).first()
            if not student:
                return jsonify({'error': 'Student record not found for current user'}), 404
            students = [student]
        else:
            # Get all students with their class information
            students = Student.query.join(Class).join(User).all()
        
        # Group students by class
        class_data = {}
        for student in students:
            class_name = student.class_.class_name
            if class_name not in class_data:
                class_data[class_name] = {
                    'class_name': class_name,
                    'students': []
                }
            
            # Get fees for this student
            fees = Fee.query.filter_by(student_id=student.student_id).first()
            
            # Organize fees with paid, remaining, and total
            student_fees = {
                'student_id': student.student_id,
                'roll_number': student.roll_no,
                'name': student.user.name,
                'total_fees': 0,
                'paid_fees': 0,
                'remaining_fees': 0,
                'payment_status': 'Unpaid'
            }
            
            if fees:
                student_fees['total_fees'] = float(fees.total_amount)
                student_fees['paid_fees'] = float(fees.paid_amount)
                student_fees['remaining_fees'] = float(fees.due_amount)
                student_fees['payment_status'] = fees.payment_status
            
            class_data[class_name]['students'].append(student_fees)
        
        return jsonify(list(class_data.values())), 200
    except Exception as e:
        return jsonify({'error': f'Failed to fetch fees data: {str(e)}'}), 500

@app.route('/api/student-services/announcements', methods=['GET'])
def get_student_services_announcements():
    """Get announcements for student services"""
    try:
        announcements = Announcement.query.filter_by(is_active=True).order_by(Announcement.created_at.desc()).all()
        return jsonify([announcement.to_dict() for announcement in announcements]), 200
    except Exception as e:
        return jsonify({'error': f'Failed to fetch announcements: {str(e)}'}), 500

@app.route('/api/student-services/events', methods=['GET'])
def get_student_services_events():
    """Get events for student services"""
    try:
        events = Event.query.filter_by(is_active=True).order_by(Event.event_date.desc()).all()
        return jsonify([event.to_dict() for event in events]), 200
    except Exception as e:
        return jsonify({'error': f'Failed to fetch events: {str(e)}'}), 500

@app.route('/api/student-services/notifications', methods=['GET'])
def get_student_services_notifications():
    """Get notifications for student services"""
    try:
        # In production, verify JWT token here
        user, error = get_current_user_from_token()
        if error:
            return error
            
        notifications = Notification.query.filter_by(user_id=user.user_id).order_by(Notification.created_at.desc()).all()
        return jsonify([notif.to_dict() for notif in notifications]), 200
            
    except Exception as e:
        return jsonify({'error': f'Failed to fetch notifications: {str(e)}'}), 500

@app.route('/api/student-services/dashboard', methods=['GET'])
def get_student_services_dashboard():
    """Get comprehensive dashboard data for student services"""
    try:
        # Get basic statistics
        total_students = Student.query.filter_by(is_active=True).count()
        total_classes = Class.query.filter_by(is_active=True).count()
        total_subjects = Subject.query.filter_by(is_active=True).count()
        
        # Get recent announcements
        recent_announcements = Announcement.query.filter_by(is_active=True).order_by(Announcement.created_at.desc()).limit(5).all()
        
        # Get upcoming events
        upcoming_events = Event.query.filter_by(is_active=True).filter(Event.event_date >= date.today()).order_by(Event.event_date.asc()).limit(5).all()
        
        # Get marks statistics
        total_marks = Mark.query.count()
        avg_marks = db.session.query(db.func.avg(Mark.obtained_marks)).scalar()
        
        # Get attendance statistics
        total_attendance_records = Attendance.query.count()
        avg_attendance = db.session.query(db.func.avg(Attendance.attendance_percentage)).scalar()
        
        # Get fees statistics
        total_fees_records = Fee.query.count()
        paid_fees = Fee.query.filter_by(payment_status='Paid').count()
        unpaid_fees = Fee.query.filter_by(payment_status='Unpaid').count()
        partial_fees = Fee.query.filter_by(payment_status='Partial').count()
        
        # Calculate class-wise stats
        class_stats = []
        classes = Class.query.filter_by(is_active=True).all()
        for c in classes:
            students = Student.query.filter_by(class_id=c.class_id, is_active=True).all()
            student_ids = [s.student_id for s in students]
            
            if not student_ids:
                continue
                
            # Class marks
            marks = Mark.query.filter(Mark.student_id.in_(student_ids)).all()
            total_obtained = sum(m.obtained_marks for m in marks)
            total_max = sum((m.total_marks or 35) for m in marks)
            marks_pct = (total_obtained / total_max * 100) if total_max > 0 else 0
            
            # Class attendance
            attendances = Attendance.query.filter(Attendance.student_id.in_(student_ids)).all()
            avg_attendance = sum(float(a.attendance_percentage) for a in attendances) / len(attendances) if attendances else 0
            
            # Class fees
            fees = Fee.query.filter(Fee.student_id.in_(student_ids)).all()
            paid_c_fees = sum(1 for f in fees if f.payment_status == 'Paid')
            fees_pct = (paid_c_fees / len(fees) * 100) if fees else 0
            
            class_stats.append({
                'class_name': c.class_name,
                'marks_pct': marks_pct,
                'attendance_pct': avg_attendance,
                'fees_pct': fees_pct
            })

        
        dashboard_data = {
            'statistics': {
                'total_students': total_students,
                'total_classes': total_classes,
                'total_subjects': total_subjects,
                'total_marks_records': total_marks,
                'average_marks': round(avg_marks, 2) if avg_marks else 0,
                'total_attendance_records': total_attendance_records,
                'average_attendance': round(float(avg_attendance), 2) if avg_attendance else 0,
                'total_fees_records': total_fees_records,
                'paid_fees': paid_fees,
                'unpaid_fees': unpaid_fees,
                'partial_fees': partial_fees
            },
            'class_stats': class_stats,
            'recent_announcements': [announcement.to_dict() for announcement in recent_announcements],
            'upcoming_events': [event.to_dict() for event in upcoming_events]
        }
        
        return jsonify(dashboard_data), 200
    except Exception as e:
        return jsonify({'error': f'Failed to fetch dashboard data: {str(e)}'}), 500

#
# Admin-only endpoints for editing student data & subject configuration
#

ALLOWED_TOTAL_VALUES = (25, 35, 50, 70, 100, 125)


@app.route('/api/admin/classes/<int:class_id>/subjects', methods=['POST'])
def admin_add_class_subject(class_id):
    """Add a new subject to a class and initialize records for students."""
    try:
        data = request.get_json() or {}
        required = ['subject_name', 'subject_code']
        for field in required:
            if not data.get(field):
                return jsonify({'error': f'{field} is required'}), 400
        
        cls = db.session.get(Class, class_id)
        if not cls: return jsonify({'error': 'Class not found'}), 404
            
        subject = Subject(
            class_id=class_id,
            subject_name=data['subject_name'],
            subject_code=data['subject_code'],
            credits=data.get('credits', 4),
            total_marks=data.get('total_marks', 35),
            total_classes=data.get('total_classes', 50),
            description=data.get('description', '')
        )
        db.session.add(subject)
        db.session.flush() 
        
        students = Student.query.filter_by(class_id=class_id, is_active=True).all()
        for student in students:
            db.session.add(Mark(student_id=student.student_id, subject_id=subject.subject_id, total_marks=subject.total_marks, obtained_marks=0))
            db.session.add(Attendance(student_id=student.student_id, subject_id=subject.subject_id, present_count=0, total_classes=subject.total_classes, attendance_percentage=0))
        
        db.session.commit()
        return jsonify({'message': 'Subject added and records initialized', 'subject': subject.to_dict()}), 201
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'Failed to add subject: {str(e)}'}), 500


@app.route('/api/admin/subjects/<int:subject_id>', methods=['DELETE'])
def admin_delete_subject(subject_id):
    """Delete a subject and all associated mark/attendance records."""
    try:
        subject = db.session.get(Subject, subject_id)
        if not subject: return jsonify({'error': 'Subject not found'}), 404
            
        Mark.query.filter_by(subject_id=subject_id).delete()
        Attendance.query.filter_by(subject_id=subject_id).delete()
        
        db.session.delete(subject)
        db.session.commit()
        return jsonify({'message': 'Subject and associated records deleted'}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'Failed to delete subject: {str(e)}'}), 500


@app.route('/api/admin/classes/<int:class_id>/subjects', methods=['GET'])
def admin_get_class_subjects(class_id):
    """List subjects for a class (for admin subject editor)."""
    try:
        cls = db.session.get(Class, class_id)
        if not cls:
            return jsonify({'error': 'Class not found'}), 404
        subjects = Subject.query.filter_by(class_id=class_id, is_active=True).order_by(Subject.subject_id).all()
        return jsonify([
            {
                'subject_id': s.subject_id,
                'subject_name': s.subject_name,
                'subject_code': s.subject_code
            } for s in subjects
        ]), 200
    except Exception as e:
        return jsonify({'error': f'Failed to fetch subjects: {str(e)}'}), 500


@app.route('/api/admin/subjects/<int:subject_id>', methods=['PUT'])
def admin_update_subject(subject_id):
    """Update subject metadata (currently name only) for admin."""
    try:
        data = request.get_json() or {}
        subject = db.session.get(Subject, subject_id)
        if not subject:
            return jsonify({'error': 'Subject not found'}), 404
        if 'subject_name' in data and data['subject_name']:
            subject.subject_name = data['subject_name']
        db.session.commit()
        return jsonify({'message': 'Subject updated', 'subject': subject.to_dict()}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'Failed to update subject: {str(e)}'}), 500


@app.route('/api/admin/classes/<int:class_id>/subjects/<int:subject_id>/total-marks', methods=['PUT'])
def admin_set_class_subject_total_marks(class_id, subject_id):
    """Set total marks for a subject for ALL students in the class in one go."""
    try:
        data = request.get_json() or {}
        total_marks = int(data.get('total_marks', 35))
        if total_marks not in ALLOWED_TOTAL_VALUES:
            return jsonify({'error': f'Total marks must be one of {list(ALLOWED_TOTAL_VALUES)}'}), 400

        subject = db.session.get(Subject, subject_id)
        if not subject: return jsonify({'error': 'Subject not found'}), 404
        
        subject.total_marks = total_marks
        
        students = Student.query.filter_by(class_id=class_id, is_active=True).all()
        updated = 0
        for student in students:
            mark = Mark.query.filter_by(student_id=student.student_id, subject_id=subject_id).first()
            if mark:
                mark.total_marks = total_marks
                mark.updated_at = datetime.utcnow()
                updated += 1
        db.session.commit()
        return jsonify({'message': f'Total marks set to {total_marks} for {updated} students'}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500


@app.route('/api/admin/classes/<int:class_id>/subjects/<int:subject_id>/total-classes', methods=['PUT'])
def admin_set_class_subject_total_classes(class_id, subject_id):
    """Set total classes for a subject for ALL students in the class (attendance)."""
    try:
        data = request.get_json() or {}
        total_classes = int(data.get('total_classes', 50))
        if total_classes not in ALLOWED_TOTAL_VALUES:
            return jsonify({'error': f'Total classes must be one of {list(ALLOWED_TOTAL_VALUES)}'}), 400

        subject = db.session.get(Subject, subject_id)
        if not subject: return jsonify({'error': 'Subject not found'}), 404
        
        subject.total_classes = total_classes
        
        students = Student.query.filter_by(class_id=class_id, is_active=True).all()
        updated = 0
        for student in students:
            attendance = Attendance.query.filter_by(student_id=student.student_id, subject_id=subject_id).first()
            if attendance:
                attendance.total_classes = total_classes
                # Recompute percentage if we have present_count
                if attendance.present_count is not None:
                    attendance.attendance_percentage = Decimal(
                        (attendance.present_count / total_classes) * 100
                    ).quantize(Decimal('0.01'))
                attendance.updated_at = datetime.utcnow()
                updated += 1
        db.session.commit()
        return jsonify({'message': f'Total classes set to {total_classes} for {updated} students'}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500


@app.route('/api/admin/students/<int:student_id>/marks', methods=['PUT'])
def admin_update_student_marks(student_id):
    try:
        # Check if user is admin (you can add proper authentication here)
        data = request.get_json()
        
        if not data or 'marks' not in data:
            return jsonify({'error': 'Marks data is required'}), 400
        
        student = db.session.get(Student, student_id)
        if not student:
            return jsonify({'error': 'Student not found'}), 404
        
        # Update marks for each subject
        for mark_data in data['marks']:
            if 'subject_id' not in mark_data or 'obtained_marks' not in mark_data:
                continue
                
            subject_id = mark_data['subject_id']
            obtained_marks = int(mark_data.get('obtained_marks', 0))
            total_marks = int(mark_data.get('total_marks', 35))
            
            if obtained_marks < 0:
                return jsonify({'error': f'Obtained marks cannot be negative for subject {subject_id}'}), 400
            if total_marks not in ALLOWED_TOTAL_VALUES:
                return jsonify({'error': f'Total marks must be one of {list(ALLOWED_TOTAL_VALUES)} for subject {subject_id}'}), 400
            if obtained_marks > total_marks:
                return jsonify({'error': f'Obtained marks cannot exceed total marks ({total_marks}) for subject {subject_id}'}), 400
            
            # Find existing mark or create new one
            mark = Mark.query.filter_by(student_id=student_id, subject_id=subject_id).first()
            if mark:
                mark.obtained_marks = obtained_marks
                mark.total_marks = total_marks
                mark.updated_at = datetime.utcnow()
            else:
                mark = Mark(
                    student_id=student_id,
                    subject_id=subject_id,
                    total_marks=total_marks,
                    obtained_marks=obtained_marks,
                    exam_date=datetime.now().date()
                )
                db.session.add(mark)
        
        db.session.commit()
        return jsonify({'message': 'Marks updated successfully'}), 200
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'Failed to update marks: {str(e)}'}), 500

@app.route('/api/admin/students/<int:student_id>/attendance', methods=['PUT'])
def admin_update_student_attendance(student_id):
    try:
        data = request.get_json()
        
        if not data or 'attendance' not in data:
            return jsonify({'error': 'Attendance data is required'}), 400
        
        student = db.session.get(Student, student_id)
        if not student:
            return jsonify({'error': 'Student not found'}), 404
        
        # Get current academic year from student's admission year or use current year
        current_year = datetime.now().year
        academic_year = f"{student.admission_year}-{student.admission_year + 1}"
        
        # Update attendance for each subject
        for attendance_data in data['attendance']:
            if 'subject_id' not in attendance_data:
                continue
                
            subject_id = attendance_data['subject_id']
            present_count = int(attendance_data.get('present_count', 0))
            total_classes = int(attendance_data.get('total_classes', 50))
            late_count = int(attendance_data.get('late_count', 0))
            
            # Validate attendance data
            if present_count < 0:
                return jsonify({'error': f'Present count cannot be negative for subject {subject_id}'}), 400
            if total_classes <= 0:
                return jsonify({'error': f'Total classes must be greater than 0 for subject {subject_id}'}), 400
            if present_count > total_classes:
                return jsonify({'error': f'Present count ({present_count}) cannot exceed total classes ({total_classes}) for subject {subject_id}'}), 400
            if late_count < 0 or late_count > present_count:
                return jsonify({'error': f'Late count must be between 0 and present count for subject {subject_id}'}), 400
            
            # Calculate absent_count
            absent_count = max(0, total_classes - present_count)
            
            # Calculate attendance percentage
            attendance_percentage = Decimal((present_count / total_classes) * 100).quantize(Decimal('0.01')) if total_classes > 0 else Decimal(0)
            
            # Find existing attendance record for this student and subject
            # If multiple exist (different academic years), update the most recent one
            attendance = Attendance.query.filter_by(
                student_id=student_id, 
                subject_id=subject_id
            ).order_by(Attendance.updated_at.desc()).first()
            
            # If still not found, try without academic year filter (in case academic_year is null)
            if not attendance:
                attendance = Attendance.query.filter_by(
                    student_id=student_id, 
                    subject_id=subject_id
                ).first()
            
            if attendance:
                # Update existing attendance record using direct assignment
                print(f"[DEBUG] Updating attendance ID {attendance.attendance_id} for student {student_id}, subject {subject_id}")
                print(f"[DEBUG] Old: present={attendance.present_count}/{attendance.total_classes} ({attendance.attendance_percentage}%)")
                print(f"[DEBUG] New: present={present_count}/{total_classes} ({attendance_percentage}%)")
                
                # Direct update - SQLAlchemy will track these changes automatically
                attendance.present_count = present_count
                attendance.absent_count = absent_count
                attendance.late_count = late_count
                attendance.total_classes = total_classes
                attendance.attendance_percentage = attendance_percentage
                attendance.academic_year = academic_year
                attendance.updated_at = datetime.utcnow()
                
                # Ensure the object is in the session (it should be, but just to be sure)
                if attendance not in db.session:
                    db.session.add(attendance)
            else:
                # Create new attendance record
                print(f"Creating new attendance record for student {student_id}, subject {subject_id}")
                attendance = Attendance(
                    student_id=student_id,
                    subject_id=subject_id,
                    present_count=present_count,
                    absent_count=absent_count,
                    late_count=late_count,
                    total_classes=total_classes,
                    attendance_percentage=attendance_percentage,
                    academic_year=academic_year,
                    updated_at=datetime.utcnow()
                )
                db.session.add(attendance)
        
        # Commit all changes with proper error handling
        verification_data = []
        try:
            # Check what's in the session before commit
            print(f"[DEBUG] Session has {len(db.session.dirty)} dirty objects before commit")
            print(f"[DEBUG] Session has {len(db.session.new)} new objects before commit")
            
            db.session.commit()
            print(f"[DEBUG] Successfully committed attendance updates for student {student_id}")
            
            # Verify the update by querying fresh from database
            db.session.expire_all()  # Clear any cached objects
            verification = Attendance.query.filter_by(student_id=student_id).all()
            print(f"[DEBUG] Verification: Found {len(verification)} attendance records after commit")
            for v in verification:
                print(f"[DEBUG]   Subject {v.subject_id}: {v.present_count}/{v.total_classes} ({v.attendance_percentage}%)")
                verification_data.append({
                    'subject_id': v.subject_id,
                    'present_count': v.present_count,
                    'total_classes': v.total_classes,
                    'percentage': float(v.attendance_percentage) if v.attendance_percentage else 0
                })
                
        except Exception as commit_error:
            print(f"[ERROR] Error committing attendance: {str(commit_error)}")
            import traceback
            traceback.print_exc()
            db.session.rollback()
            raise
        
        # Return success with verification data
        return jsonify({
            'message': 'Attendance updated successfully',
            'updated_records': len(verification_data),
            'verification': verification_data
        }), 200
        
    except ValueError as e:
        db.session.rollback()
        return jsonify({'error': f'Invalid data format: {str(e)}'}), 400
    except Exception as e:
        db.session.rollback()
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'Failed to update attendance: {str(e)}'}), 500

@app.route('/api/admin/students/<int:student_id>/fees', methods=['PUT'])
def admin_update_student_fees(student_id):
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({'error': 'Fees data is required'}), 400
        
        student = db.session.get(Student, student_id)
        if not student:
            return jsonify({'error': 'Student not found'}), 404
        
        total_amount = data.get('total_amount', 0)
        paid_amount = data.get('paid_amount', 0)
        
        if paid_amount > total_amount:
            return jsonify({'error': 'Paid amount cannot exceed total amount'}), 400
        
        due_amount = total_amount - paid_amount
        
        if paid_amount == total_amount:
            payment_status = 'Paid'
        elif paid_amount == 0:
            payment_status = 'Unpaid'
        else:
            payment_status = 'Partial'
        
        # Find existing fee or create new one
        fee = Fee.query.filter_by(student_id=student_id).first()
        if fee:
            fee.total_amount = total_amount
            fee.paid_amount = paid_amount
            fee.due_amount = due_amount
            fee.payment_status = payment_status
            fee.updated_at = datetime.utcnow()
        else:
            fee = Fee(
                student_id=student_id,
                total_amount=total_amount,
                paid_amount=paid_amount,
                due_amount=due_amount,
                payment_status=payment_status,
                last_payment_date=datetime.now().date() if paid_amount > 0 else None,
                created_at=datetime.utcnow(),
                updated_at=datetime.utcnow()
            )
            db.session.add(fee)
        
        db.session.commit()
        return jsonify({'message': 'Fees updated successfully'}), 200
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'Failed to update fees: {str(e)}'}), 500

@app.route('/api/admin/students/<int:student_id>/data', methods=['GET'])
def admin_get_student_data(student_id):
    try:
        student = db.session.get(Student, student_id)
        if not student:
            return jsonify({'error': 'Student not found'}), 404
        
        # Get student basic info
        student_data = {
            'student_id': student.student_id,
            'roll_no': student.roll_no,
            'name': student.user.name,
            'class_name': student.class_.class_name,
            'marks': [],
            'attendance': [],
            'fees': {}
        }
        
        # Always return rows for each subject in the student's class (default 0 if missing)
        subjects = Subject.query.filter_by(class_id=student.class_id).all()

        marks_rows = Mark.query.filter_by(student_id=student_id).all()
        marks_map = {m.subject_id: m for m in marks_rows}
        for subject in subjects:
            m = marks_map.get(subject.subject_id)
            # Use Subject-level total as official, but keep Mark record synced
            s_total = subject.total_marks or 35
            student_data['marks'].append({
                'subject_id': subject.subject_id,
                'subject_name': subject.subject_name,
                'obtained_marks': int(getattr(m, 'obtained_marks', 0) or 0),
                'total_marks': s_total,
                'exam_date': m.exam_date.isoformat() if m and m.exam_date else None
            })

        attendance_rows = Attendance.query.filter_by(student_id=student_id).all()
        attendance_map = {a.subject_id: a for a in attendance_rows}
        for subject in subjects:
            a = attendance_map.get(subject.subject_id)
            s_total_classes = subject.total_classes or 50
            present = int(getattr(a, 'present_count', 0) or 0)
            percent = float(getattr(a, 'attendance_percentage', 0) or 0)
            if s_total_classes > 0 and (a is None or a.attendance_percentage is None):
                percent = (present / s_total_classes) * 100
            student_data['attendance'].append({
                'subject_id': subject.subject_id,
                'subject_name': subject.subject_name,
                'present_count': present,
                'total_classes': s_total_classes,
                'attendance_percentage': percent
            })

        # Fees: always return an object (0 defaults if missing)
        fee = Fee.query.filter_by(student_id=student_id).first()
        if fee:
            student_data['fees'] = {
                'total_amount': float(fee.total_amount or 0),
                'paid_amount': float(fee.paid_amount or 0),
                'due_amount': float(fee.due_amount or 0),
                'payment_status': fee.payment_status or 'Unpaid',
                'last_payment_date': fee.last_payment_date.isoformat() if fee.last_payment_date else None
            }
        else:
            student_data['fees'] = {
                'total_amount': 0,
                'paid_amount': 0,
                'due_amount': 0,
                'payment_status': 'Unpaid',
                'last_payment_date': None
            }
        
        return jsonify(student_data), 200
        
    except Exception as e:
        return jsonify({'error': f'Failed to fetch student data: {str(e)}'}), 500

# Event management endpoints
@app.route('/api/events', methods=['POST'])
def create_event():
    try:
        data = request.get_json()
        
        if not data or 'title' not in data:
            return jsonify({'error': 'Event title is required'}), 400
        
        event_type = data.get('event_type', 'general')
        
        # Get the appropriate Google Form link based on event type
        form_links = {
            'workshop': 'https://forms.gle/ekED5EhxvY7xRjok6',
            'seminar': 'https://forms.gle/ekED5EhxvY7xRjok6',
            'hackathon': 'https://forms.gle/yCUZgrkt9hrG5m5e7',
            'club_event': 'https://forms.gle/zYYxKZkUQsv1aj4U8',
            'competition': 'https://forms.gle/jXupRmdY2Q4Hwcjr6',
            'conference': 'https://forms.gle/t5pJb3FZsWcSdDBL8',
            'cultural': 'https://forms.gle/t5pJb3FZsWcSdDBL8',
            'sports': 'https://forms.gle/t5pJb3FZsWcSdDBL8',
            'academic': 'https://forms.gle/t5pJb3FZsWcSdDBL8',
            'general': 'https://forms.gle/t5pJb3FZsWcSdDBL8'
        }
        
        # Use provided registration_link or default to Google Form
        registration_link = data.get('registration_link', '') or form_links.get(event_type, form_links['general'])
        
        event = Event(
            title=data['title'],
            description=data.get('description', ''),
            event_date=datetime.strptime(data.get('event_date', datetime.now().strftime('%Y-%m-%d')), '%Y-%m-%d').date(),
            event_time=datetime.strptime(data.get('event_time', '10:00'), '%H:%M').time() if data.get('event_time') else None,
            location=data.get('location', ''),
            event_type=event_type,
            organized_by=data.get('organized_by', ''),
            registration_link=registration_link,
            is_active=True,
            created_at=datetime.utcnow()
        )
        
        db.session.add(event)
        db.session.commit()
        
        return jsonify(event.to_dict()), 201
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'Failed to create event: {str(e)}'}), 500

@app.route('/api/events/<int:event_id>', methods=['DELETE'])
def delete_event(event_id):
    try:
        event = db.session.get(Event, event_id)
        if not event:
            return jsonify({'error': 'Event not found'}), 404
        
        db.session.delete(event)
        db.session.commit()
        
        return jsonify({'message': 'Event deleted successfully'}), 200
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'Failed to delete event: {str(e)}'}), 500

# Event Registration Endpoints
@app.route('/api/events', methods=['GET'])
def get_all_events():
    """Get all events with registration counts"""
    try:
        events = Event.query.filter_by(is_active=True).order_by(Event.event_date.desc()).all()
        
        # Get registration counts for each event
        events_data = []
        for event in events:
            event_dict = event.to_dict()
            events_data.append(event_dict)
        
        return jsonify({'success': True, 'data': events_data}), 200
    except Exception as e:
        return jsonify({'error': f'Failed to fetch events: {str(e)}'}), 500

@app.route('/api/events/<int:event_id>/register', methods=['POST'])
def register_for_event(event_id):
    """Register a student for an event"""
    try:
        # Get current user from token (you'll need to implement this based on your auth system)
        # For now, assuming user_id is passed in the request
        data = request.get_json()
        user_id = data.get('user_id') or request.headers.get('user_id')
        
        if not user_id:
            return jsonify({'error': 'User ID is required'}), 400
        
        # Get student by user_id
        student = Student.query.filter_by(user_id=user_id).first()
        if not student:
            return jsonify({'error': 'Student not found'}), 404
        
        # Check if event exists
        event = db.session.get(Event, event_id)
        if not event:
            return jsonify({'error': 'Event not found'}), 404
        
        # Check if already registered
        existing_registration = EventRegistration.query.filter_by(
            event_id=event_id,
            student_id=student.student_id
        ).first()
        
        if existing_registration:
            return jsonify({'error': 'Already registered for this event'}), 400
        
        # Check if event has reached max participants
        if event.max_participants:
            registration_count = EventRegistration.query.filter_by(event_id=event_id).count()
            if registration_count >= event.max_participants:
                return jsonify({'error': 'Event is full'}), 400
        
        # Create registration
        registration = EventRegistration(
            event_id=event_id,
            student_id=student.student_id,
            registered_at=datetime.utcnow()
        )
        
        db.session.add(registration)
        
        # Update event current participants
        event.current_participants = EventRegistration.query.filter_by(event_id=event_id).count() + 1
        
        db.session.commit()
        
        return jsonify({'success': True, 'data': registration.to_dict()}), 201
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'Failed to register for event: {str(e)}'}), 500

@app.route('/api/events/<int:event_id>/register', methods=['DELETE'])
def cancel_event_registration(event_id):
    """Cancel event registration"""
    try:
        data = request.get_json()
        user_id = data.get('user_id') or request.headers.get('user_id')
        
        if not user_id:
            return jsonify({'error': 'User ID is required'}), 400
        
        # Get student by user_id
        student = Student.query.filter_by(user_id=user_id).first()
        if not student:
            return jsonify({'error': 'Student not found'}), 404
        
        # Find registration
        registration = EventRegistration.query.filter_by(
            event_id=event_id,
            student_id=student.student_id
        ).first()
        
        if not registration:
            return jsonify({'error': 'Registration not found'}), 404
        
        # Delete registration
        db.session.delete(registration)
        
        # Update event current participants
        event = db.session.get(Event, event_id)
        if event:
            event.current_participants = EventRegistration.query.filter_by(event_id=event_id).count()
        
        db.session.commit()
        
        return jsonify({'success': True, 'message': 'Registration cancelled successfully'}), 200
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'Failed to cancel registration: {str(e)}'}), 500

@app.route('/api/events/my-registrations', methods=['GET'])
def get_my_registrations():
    """Get current user's event registrations"""
    try:
        user_id = request.headers.get('user_id') or request.args.get('user_id')
        
        if not user_id:
            return jsonify({'error': 'User ID is required'}), 400
        
        # Get student by user_id
        student = Student.query.filter_by(user_id=user_id).first()
        if not student:
            return jsonify({'success': True, 'data': []}), 200
        
        # Get all registrations for this student
        registrations = EventRegistration.query.filter_by(student_id=student.student_id).all()
        
        # Get event details for each registration
        registrations_data = []
        for registration in registrations:
            event = db.session.get(Event, registration.event_id)
            if event:
                registrations_data.append({
                    'id': registration.registration_id,
                    'event_id': registration.event_id,
                    'title': event.title,
                    'event_date': event.event_date.isoformat(),
                    'event_time': event.event_time.isoformat() if event.event_time else None,
                    'location': event.location,
                    'registered_at': registration.registered_at.isoformat()
                })
        
        return jsonify({'success': True, 'data': registrations_data}), 200
        
    except Exception as e:
        return jsonify({'error': f'Failed to fetch registrations: {str(e)}'}), 500

# Notification management endpoints
@app.route('/api/notifications', methods=['POST'])
def create_notification():
    try:
        data = request.get_json()
        
        if not data or 'title' not in data or 'message' not in data:
            return jsonify({'error': 'Notification title and message are required'}), 400
        
        # Create announcement instead of notification
        priority_map = {
            'info': 'normal',
            'warning': 'high', 
            'success': 'normal',
            'error': 'urgent'
        }
        
        announcement = Announcement(
            title=data['title'],
            message=data['message'],
            priority=priority_map.get(data.get('type', 'info'), 'normal'),
            target='all',
            is_active=True,
            created_at=datetime.utcnow()
        )
        
        db.session.add(announcement)
        db.session.commit()
        
        return jsonify(announcement.to_dict()), 201
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'Failed to create notification: {str(e)}'}), 500

@app.route('/api/notifications/<int:notification_id>', methods=['DELETE'])
def delete_notification(notification_id):
    try:
        # Delete from announcements table instead
        announcement = db.session.get(Announcement, notification_id)
        if not announcement:
            return jsonify({'error': 'Notification not found'}), 404
        
        db.session.delete(announcement)
        db.session.commit()
        
        return jsonify({'message': 'Notification deleted successfully'}), 200
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'Failed to delete notification: {str(e)}'}), 500


# HOD Management Endpoints
@app.route('/api/hod/students/<branch>/<year>', methods=['GET'])
def get_hod_students(branch, year):
    """Get students for HOD by branch and year"""
    try:
        # Find the department (by name or code). If not found, return empty list instead of 404
        dept = Department.query.filter_by(dept_name=branch).first()
        if not dept:
            dept = Department.query.filter_by(dept_code=branch).first()
        if not dept:
            return jsonify([]), 200
        
        # Derive a stable class code like "SY-CSE" and try both name and code
        class_code = f"{year}-{branch}"
        class_obj = Class.query.filter_by(class_code=class_code, dept_id=dept.dept_id).first()
        if not class_obj:
            # Fallback to class_name if older data uses it
            class_name = f"{year}-{branch}"
            class_obj = Class.query.filter_by(class_name=class_name, dept_id=dept.dept_id).first()
        if not class_obj:
            return jsonify([]), 200
        
        students = Student.query.join(User).filter(Student.class_id == class_obj.class_id).all()
        
        student_data = []
        for student in students:
            student_info = {
                'roll_no': student.roll_no,
                'name': student.user.name,
                'email': student.user.email,
                'contact': student.user.contact_no,
                'username': student.user.email.split('@')[0],  # Use email prefix as username
                'admission_year': student.admission_year
            }
            student_data.append(student_info)
        
        return jsonify(student_data), 200
    except Exception as e:
        return jsonify({'error': f'Failed to fetch students: {str(e)}'}), 500

@app.route('/api/hod/students/<branch>/<year>', methods=['POST'])
def create_hod_student(branch, year):
    """Create a new student for HOD"""
    try:
        data = request.get_json()
        
        if not data or not data.get('roll_no') or not data.get('name') or not data.get('email'):
            return jsonify({'error': 'Roll number, name, and email are required'}), 400
        
        # Find or create the department
        dept = Department.query.filter_by(dept_name=branch).first()
        if not dept:
            dept = Department.query.filter_by(dept_code=branch).first()
        if not dept:
            dept = Department(
                dept_name=f"{branch} Department",
                dept_code=branch,
                description=f"{branch} Department",
                is_active=True,
                created_at=datetime.utcnow()
            )
            db.session.add(dept)
            db.session.flush()
        
        # Find or create the class for this branch/year
        class_code = f"{year}-{branch}"
        class_obj = Class.query.filter_by(class_code=class_code, dept_id=dept.dept_id).first()
        if not class_obj:
            class_name = f"{year} {branch}"
            class_obj = Class(
                dept_id=dept.dept_id,
                class_name=class_name,
                class_code=class_code,
                academic_year=str(datetime.utcnow().year),
                is_active=True,
                created_at=datetime.utcnow()
            )
            db.session.add(class_obj)
            db.session.flush()
        
        # Check if roll number already exists in this class only (same roll allowed in SY, TY, Final)
        if Student.query.filter_by(roll_no=data['roll_no'], class_id=class_obj.class_id).first():
            return jsonify({'error': f'Roll number {data["roll_no"]} already exists in this class'}), 400
        
        # Check if email already exists
        if User.query.filter_by(email=data['email']).first():
            return jsonify({'error': 'Email already exists'}), 400
        
        # Create user first - default password equals contact number for initial login
        default_password = data.get('contact') or data.get('roll_no')
        user = User(
            name=data['name'],
            email=data['email'],
            contact_no=data.get('contact', ''),
            role='student',
            is_active=True
        )
        user.set_password(default_password or 'changeme123')
        
        db.session.add(user)
        db.session.flush()  # Get the user_id
        
        # Create student
        student = Student(
            user_id=user.user_id,
            roll_no=data['roll_no'],
            class_id=class_obj.class_id,
            admission_year=data.get('admission_year', datetime.now().year),
            admission_date=datetime.now().date(),
            guardian_name=data.get('guardian_name'),
            guardian_contact=data.get('guardian_contact'),
            address=data.get('address')
        )
        
        db.session.add(student)
        db.session.flush()  # Get the student_id
        
        # Ensure class has subjects so we can seed marks/attendance
        subjects = Subject.query.filter_by(class_id=class_obj.class_id).all()
        if not subjects:
            for i in range(1, 6):
                db.session.add(Subject(
                    class_id=class_obj.class_id,
                    subject_name=f"Subject {i}",
                    subject_code=f"{class_obj.class_code}-S{i}"[:20],
                    credits=4,
                    description="",
                    is_active=True,
                    created_at=datetime.utcnow()
                ))
            db.session.flush()
            subjects = Subject.query.filter_by(class_id=class_obj.class_id).all()

        # Create default marks for all subjects in the class using synchronized totals
        for subject in subjects:
            mark = Mark(
                student_id=student.student_id,
                subject_id=subject.subject_id,
                total_marks=subject.total_marks or 35, # Default to subject's total
                obtained_marks=0,
                exam_date=datetime.now().date()
            )
            db.session.add(mark)
        
        # Create default attendance for all subjects using synchronized totals
        for subject in subjects:
            attendance = Attendance(
                student_id=student.student_id,
                subject_id=subject.subject_id,
                present_count=0,
                absent_count=0,
                late_count=0,
                total_classes=subject.total_classes or 50, # Default to subject's total
                attendance_percentage=0.0,
                academic_year=str(datetime.now().year)
            )
            db.session.add(attendance)
        
        # Create default fees
        fee = Fee(
            student_id=student.student_id,
            total_amount=50000.00,  # Default total fees
            paid_amount=0.00,
            due_amount=50000.00,
            payment_status='Unpaid'
        )
        db.session.add(fee)
        
        db.session.commit()
        
        return jsonify({
            'roll_no': student.roll_no,
            'name': user.name,
            'email': user.email,
            'contact': user.contact_no,
            'username': user.email.split('@')[0],
            'admission_year': student.admission_year
        }), 201
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'Failed to create student: {str(e)}'}), 500

@app.route('/api/hod/students/<branch>/<year>/<roll_no>', methods=['PUT'])
def update_hod_student(branch, year, roll_no):
    """Update a student for HOD"""
    try:
        data = request.get_json()
        
        # Find the student
        student = Student.query.filter_by(roll_no=roll_no).first()
        if not student:
            return jsonify({'error': 'Student not found'}), 404
        
        # Update user information
        user = student.user
        if data.get('name'):
            user.name = data['name']
        if data.get('email'):
            user.email = data['email']
        if data.get('contact'):
            user.contact_no = data['contact']
        if data.get('admission_year'):
            student.admission_year = data['admission_year']
        
        db.session.commit()
        
        return jsonify({
            'roll_no': student.roll_no,
            'name': user.name,
            'email': user.email,
            'contact': user.contact_no,
            'username': user.email.split('@')[0],
            'admission_year': student.admission_year
        }), 200
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'Failed to update student: {str(e)}'}), 500

@app.route('/api/hod/students/<branch>/<year>/<roll_no>', methods=['DELETE'])
def delete_hod_student(branch, year, roll_no):
    """Delete a student for HOD"""
    try:
        # Find the student
        student = Student.query.filter_by(roll_no=roll_no).first()
        if not student:
            return jsonify({'error': 'Student not found'}), 404
        
        # Delete related data first
        Mark.query.filter_by(student_id=student.student_id).delete()
        Attendance.query.filter_by(student_id=student.student_id).delete()
        Fee.query.filter_by(student_id=student.student_id).delete()
        
        # Delete student and user
        user_id = student.user_id
        db.session.delete(student)
        User.query.filter_by(user_id=user_id).delete()
        
        db.session.commit()
        
        return jsonify({'message': 'Student deleted successfully'}), 200
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'Failed to delete student: {str(e)}'}), 500


# ============ AI SERVICES HELPER FUNCTIONS ============
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'docx', 'doc'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_file(file):
    filename = secure_filename(file.filename)
    ext = filename.rsplit('.', 1)[1].lower()
    try:
        if ext == "pdf":
            reader = PyPDF2.PdfReader(file)
            return "\n".join([page.extract_text() or "" for page in reader.pages])
        elif ext in ["docx", "doc"]:
            document = docx.Document(file)
            return "\n".join([p.text for p in document.paragraphs])
        elif ext == "txt":
            return file.read().decode('utf-8')
    except Exception as e:
        print("File extraction error:", e)
        return None

def clean_text(text):
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

def generate_notes(content, num_sentences=12):
    content = clean_text(content)
    if not content:
        return {"error": "No content provided."}

    # Simple note generation: extract sentences and create bullet points
    sentences = content.split('.')
    notes = []
    
    # Get key sentences (first N sentences)
    key_sentences = [s.strip() for s in sentences if len(s.strip()) > 20][:num_sentences]
    
    for i, sentence in enumerate(key_sentences):
        if sentence:
            notes.append(f"{i+1}. {sentence}")

    return {"notes": "\n".join(notes)}

@app.route('/api/ai/health', methods=['GET'])
def ai_health():
    """Health check for AI services"""
    return jsonify({"ok": True}), 200

# Redundant admission endpoints removed
# ============ PROFILE UPDATE API ============
@app.route('/api/users/<int:user_id>', methods=['PUT'])
def update_user_profile(user_id):
    """Update user profile"""
    try:
        data = request.get_json()
        
        user = db.session.get(User, user_id)
        if not user:
            return jsonify({'error': 'User not found'}), 404
        
        # Update allowed fields
        if 'name' in data:
            user.name = data['name']
        if 'email' in data:
            user.email = data['email']
        if 'contact_no' in data:
            user.contact_no = data['contact_no']
        
        # Update password if provided
        if 'password' in data and data['password']:
            # In a real application, hash the password before storing
            user.password = data['password']
        
        db.session.commit()
        
        return jsonify({'success': True, 'message': 'Profile updated successfully', 'user': user.to_dict()}), 200
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'Failed to update profile: {str(e)}'}), 500

# ============ FINE MANAGEMENT API ============
@app.route('/api/fines', methods=['GET'])
def get_all_fines():
    """Get all fines (admin only) or fines for a specific student"""
    try:
        user, error = get_current_user_from_token()
        if error:
            return error

        student_id = request.args.get('student_id', type=int)
        
        # If student, force student_id to their own
        if user.role == 'student':
            student = Student.query.filter_by(user_id=user.user_id).first()
            if not student:
                return jsonify({'error': 'Student record not found'}), 404
            student_id = student.student_id

        query = Fine.query
        if student_id:
            query = query.filter_by(student_id=student_id)
            
        # Order by newest first
        query = query.order_by(Fine.issued_date.desc())
        
        fines = query.all()
        return jsonify([fine.to_dict() for fine in fines]), 200
    except Exception as e:
        return jsonify({'error': f'Failed to fetch fines: {str(e)}'}), 500

@app.route('/api/fines', methods=['POST'])
def issue_fine():
    """Issue a new fine (admin only)"""
    try:
        user, error = get_current_user_from_token()
        if error:
            return error
            
        if user.role == 'student':
            return jsonify({'error': 'Not authorized to issue fines'}), 403
            
        data = request.get_json()
        required_fields = ['student_id', 'reason', 'amount']
        
        for field in required_fields:
            if not data.get(field):
                return jsonify({'error': f'{field} is required'}), 400
                
        # Verify student exists
        student = db.session.get(Student, data['student_id'])
        if not student:
            return jsonify({'error': 'Student not found'}), 404
            
        fine = Fine(
            student_id=data['student_id'],
            reason=data['reason'],
            amount=data['amount'],
            status='Pending'
        )
        
        db.session.add(fine)
        db.session.commit()
        
        return jsonify(fine.to_dict()), 201
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'Failed to issue fine: {str(e)}'}), 500

@app.route('/api/fines/<int:fine_id>/pay', methods=['POST'])
def pay_fine(fine_id):
    """Submit a receipt for a fine (student)"""
    try:
        user, error = get_current_user_from_token()
        if error:
            return error
            
        data = request.get_json()
        if not data or not data.get('receipt_data'):
            return jsonify({'error': 'Receipt image (base64) is required'}), 400
            
        fine = db.session.get(Fine, fine_id)
        if not fine:
            return jsonify({'error': 'Fine not found'}), 404
            
        # Verify ownership if student
        if user.role == 'student':
            student = Student.query.filter_by(user_id=user.user_id).first()
            if not student or fine.student_id != student.student_id:
                return jsonify({'error': 'Not authorized to pay this fine'}), 403
                
        # Update fine status
        fine.receipt_data = data['receipt_data']
        fine.status = 'Pending Approval'
        
        db.session.commit()
        return jsonify(fine.to_dict()), 200
    except Exception as e:
        import traceback
        with open('error_log.txt', 'w') as f:
            f.write(traceback.format_exc())
            f.write(f"\nData size: len {len(str(data))} bytes\n")
        db.session.rollback()
        return jsonify({'error': f'Failed to submit payment: {str(e)}'}), 500

@app.route('/api/fines/<int:fine_id>/approve', methods=['POST'])
def approve_fine(fine_id):
    """Approve a fine payment (admin only)"""
    try:
        user, error = get_current_user_from_token()
        if error:
            return error
            
        if user.role == 'student':
            return jsonify({'error': 'Not authorized to approve fines'}), 403
            
        fine = db.session.get(Fine, fine_id)
        if not fine:
            return jsonify({'error': 'Fine not found'}), 404
            
        # Update fine status
        fine.status = 'Approved'
        
        db.session.commit()
        
        return jsonify(fine.to_dict()), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'Failed to approve fine: {str(e)}'}), 500

@app.route('/api/fines/<int:fine_id>', methods=['DELETE'])
def delete_fine(fine_id):
    """Delete a fine (admin only)"""
    try:
        user, error = get_current_user_from_token()
        if error:
            return error
            
        if user.role == 'student':
            return jsonify({'error': 'Not authorized to delete fines'}), 403
            
        fine = db.session.get(Fine, fine_id)
        if not fine:
            return jsonify({'error': 'Fine not found'}), 404
            
        db.session.delete(fine)
        db.session.commit()
        
        return jsonify({'message': 'Fine deleted successfully'}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'Failed to delete fine: {str(e)}'}), 500


# ============================================================
# NEW ENDPOINTS — Feature Parity for Final Year Project
# ============================================================

# --- 1. Chat History ---
@app.route('/chatbot/history', methods=['GET'])
def chatbot_history():
    """Return the last 20 chat messages for the authenticated user."""
    user, err = get_current_user_from_token()
    if err:
        return err
    try:
        messages = ChatMessage.query.filter_by(user_id=user.user_id)\
            .order_by(ChatMessage.timestamp.desc()).limit(20).all()
        return jsonify([m.to_dict() for m in reversed(messages)])
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# --- 2. Proactive Alerts ---
@app.route('/chatbot/alerts', methods=['GET'])
def chatbot_alerts():
    """Return proactive academic alerts for the current student user."""
    user, err = get_current_user_from_token()
    if err:
        return err

    alerts = []
    try:
        if user.role == 'student':
            student = Student.query.filter_by(user_id=user.user_id).first()
            if student:
                # Check low attendance
                from sqlalchemy import func
                att_data = db.session.query(
                    func.sum(Attendance.present_count),
                    func.sum(Attendance.total_classes)
                ).filter_by(student_id=student.student_id).first()

                if att_data and att_data[1] and att_data[1] > 0:
                    pct = (att_data[0] / att_data[1]) * 100
                    if pct < 75:
                        alerts.append({
                            'type': 'warning',
                            'icon': '⚠️',
                            'title': 'Low Attendance',
                            'message': f'Your overall attendance is {pct:.1f}%, which is below the required 75%. Please attend classes regularly to avoid exam debarment.'
                        })

                # Check low marks (< 35%)
                low_marks = db.session.query(Mark, Subject).join(
                    Subject, Mark.subject_id == Subject.subject_id
                ).filter(
                    Mark.student_id == student.student_id
                ).all()

                for mark, subject in low_marks:
                    if mark.total_marks and mark.total_marks > 0:
                        pct = (float(mark.obtained_marks) / float(mark.total_marks)) * 100
                        if pct < 35:
                            alerts.append({
                                'type': 'danger',
                                'icon': '📉',
                                'title': f'Low Marks — {subject.subject_name}',
                                'message': f'You scored {mark.obtained_marks}/{mark.total_marks} ({pct:.1f}%) in {subject.subject_name}. Consider reviewing study materials or seeking help.'
                            })

                # Check pending fees
                fee = Fee.query.filter_by(student_id=student.student_id).first()
                if fee and fee.due_amount and float(fee.due_amount) > 0:
                    alerts.append({
                        'type': 'info',
                        'icon': '💳',
                        'title': 'Pending Fee Payment',
                        'message': f'You have ₹{float(fee.due_amount):,.0f} in pending fees. Please clear your dues before the deadline.'
                    })

                # Check upcoming events
                upcoming = Event.query.filter(
                    Event.is_active == True,
                    Event.event_date >= date.today()
                ).order_by(Event.event_date.asc()).limit(2).all()
                for ev in upcoming:
                    alerts.append({
                        'type': 'success',
                        'icon': '🎉',
                        'title': f'Upcoming: {ev.title}',
                        'message': f'Event on {ev.event_date}. Click Events to register!'
                    })

    except Exception as e:
        print(f"Alerts error: {e}")

    return jsonify({'alerts': alerts})


# --- 3. Latest Notifications (for polling) ---
@app.route('/api/notifications/latest', methods=['GET'])
def latest_notifications():
    """Return latest notifications + announcements for the logged-in user."""
    user, err = get_current_user_from_token()
    if err:
        return err
    try:
        # Get unread count from user notifications
        unread_count = Notification.query.filter_by(
            user_id=user.user_id, is_read=False
        ).count()

        # Get latest 5 system-wide announcements
        announcements = Announcement.query.filter_by(is_active=True)\
            .order_by(Announcement.created_at.desc()).limit(5).all()

        return jsonify({
            'unread_count': unread_count,
            'announcements': [a.to_dict() for a in announcements]
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# --- 4. Admin Analytics: Chatbot Insights ---
@app.route('/api/analytics/chatbot', methods=['GET'])
def analytics_chatbot():
    """Return chatbot usage analytics for admin dashboard."""
    user, err = get_current_user_from_token()
    if err:
        return err
    if user.role != 'admin':
        return jsonify({'error': 'Admin access required'}), 403

    try:
        from sqlalchemy import func

        # Intent frequency (top 10)
        intent_data = db.session.query(
            ChatMessage.intent,
            func.count(ChatMessage.id).label('count')
        ).filter(ChatMessage.intent != None)\
         .group_by(ChatMessage.intent)\
         .order_by(func.count(ChatMessage.id).desc())\
         .limit(10).all()

        # Sentiment distribution
        sentiment_data = db.session.query(
            ChatMessage.sentiment_label,
            func.count(ChatMessage.id).label('count')
        ).filter(ChatMessage.sentiment_label != None)\
         .group_by(ChatMessage.sentiment_label).all()

        # Daily message volume (last 7 days)
        from datetime import timedelta
        today = datetime.utcnow().date()
        daily_data = []
        for i in range(6, -1, -1):
            day = today - timedelta(days=i)
            count = ChatMessage.query.filter(
                func.date(ChatMessage.timestamp) == day
            ).count()
            daily_data.append({'date': day.strftime('%b %d'), 'messages': count})

        # Total stats
        total_messages = ChatMessage.query.count()
        total_users = User.query.count()

        return jsonify({
            'total_messages': total_messages,
            'total_users': total_users,
            'intent_frequency': [{'intent': r[0] or 'unknown', 'count': r[1]} for r in intent_data],
            'sentiment_distribution': [{'label': r[0] or 'neutral', 'count': r[1]} for r in sentiment_data],
            'daily_volume': daily_data
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# --- 5. Admin Analytics: Student Performance ---
@app.route('/api/analytics/students', methods=['GET'])
def analytics_students():
    """Return student academic analytics for admin dashboard."""
    user, err = get_current_user_from_token()
    if err:
        return err
    if user.role != 'admin':
        return jsonify({'error': 'Admin access required'}), 403

    try:
        from sqlalchemy import func

        # Average marks per subject
        marks_data = db.session.query(
            Subject.subject_name,
            func.avg(Mark.obtained_marks / Mark.total_marks * 100).label('avg_pct')
        ).join(Mark, Mark.subject_id == Subject.subject_id)\
         .filter(Mark.total_marks > 0)\
         .group_by(Subject.subject_name)\
         .order_by(func.avg(Mark.obtained_marks / Mark.total_marks * 100).asc())\
         .limit(8).all()

        # Attendance distribution per class
        att_data = db.session.query(
            Class.class_name,
            func.round(func.avg(
                Attendance.present_count * 100.0 / Attendance.total_classes
            ), 1).label('avg_attendance')
        ).join(Student, Student.class_id == Class.class_id)\
         .join(Attendance, Attendance.student_id == Student.student_id)\
         .filter(Attendance.total_classes > 0)\
         .group_by(Class.class_name).all()

        # Fee payment summary
        fee_status = db.session.query(
            Fee.payment_status,
            func.count(Fee.fee_id).label('count')
        ).group_by(Fee.payment_status).all()

        return jsonify({
            'subject_performance': [{'subject': r[0], 'avg_percentage': round(float(r[1] or 0), 1)} for r in marks_data],
            'class_attendance': [{'class_name': r[0], 'avg_attendance': float(r[1] or 0)} for r in att_data],
            'fee_status': [{'status': r[0] or 'unknown', 'count': r[1]} for r in fee_status]
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500



# ============================================================
# ADMISSION INFO ENDPOINTS — Served directly from JSON files
# ============================================================

def _load_json_data(filename):
    """Load a JSON file from the project's data directory."""
    data_dir = os.path.join(os.path.dirname(__file__), '..', 'data')
    try:
        with open(os.path.join(data_dir, filename), 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        print(f"Error loading {filename}: {e}")
        return {}


@app.route('/api/admission/info', methods=['GET'])
def admission_info():
    """Return admission info from Admission.json, Institute Info.json, Placements.json."""
    admission = _load_json_data('Admission.json')
    institute = _load_json_data('Institute Info.json')
    placements = _load_json_data('Placements.json')

    # Build UG branches list from JSON
    branches = []
    for course in admission.get('ug_courses', []):
        branches.append({
            'code': course.get('choice_code', ''),
            'name': course.get('course_name', ''),
            'intake': course.get('intake', 0),
            'duration': '4 Years',
            'entranceExam': 'MHT-CET / JEE Main',
            'tfws_code': course.get('tfws_choice_code', ''),
            'description': f"Undergraduate B.Tech programme with {course.get('intake', 0)} seats."
        })

    # Build admission process steps
    process_steps = [
        {'step': 1, 'title': 'Appear in Entrance Exam', 'description': 'Appear in MHT-CET / JEE Main for UG admissions.', 'deadline': 'As per DTE Maharashtra schedule'},
        {'step': 2, 'title': 'Register on DTE Portal', 'description': 'Register on the official DTE Maharashtra portal (dtemaharashtra.gov.in).', 'deadline': 'CAP Round Registration'},
        {'step': 3, 'title': 'Fill Choice Form', 'description': 'Fill preference form and choose SKN SKNSCOE – CSE / MECH / CIVIL / E&TC / EE / AI&DS.', 'deadline': 'During CAP Round'},
        {'step': 4, 'title': 'Check Allotment', 'description': 'Check seat allotment result. If college is allotted, confirm by paying seat acceptance fee.', 'deadline': 'Post CAP Round'},
        {'step': 5, 'title': 'Document Verification', 'description': 'Submit original documents at the college admission office.', 'deadline': 'Within 3 days of allotment'},
        {'step': 6, 'title': 'Fee Payment', 'description': 'Pay the prescribed annual tuition fees to confirm enrollment.', 'deadline': 'On reporting date'},
    ]

    # Required documents
    required_docs = [
        'HSC (12th) Marksheet & Certificate',
        'SSC (10th) Marksheet & Certificate',
        'MHT-CET / JEE Score Card',
        'Transfer Certificate (TC) from previous institution',
        'Migration Certificate (if from other board)',
        'Caste Certificate (if applicable)',
        'Non-Creamy Layer Certificate (if applicable)',
        'EBC/EWS Certificate (if applicable)',
        'Aadhaar Card',
        '4 Recent Passport Size Photographs',
        'Gap Certificate (if applicable)',
        'Medical Certificate (fitness)',
        'DTE Allotment Letter',
    ]

    # Placement stats mapping from Placements.json
    placement_data = placements.get('placement_statistics', {})
    top_packages = placement_data.get('top_packages', [])
    highest = f"{top_packages[0].get('package_lpa', '10.0')} LPA" if top_packages else "10.0 LPA"
    
    # Institute contact
    location = institute.get('location', {})
    
    data = {
        'college_name': institute.get('college_name', 'SKN SINHGAD COLLEGE OF ENGINEERING'),
        'established_year': institute.get('about', {}).get('establishment_date', '2010'),
        'campus_area_acres': institute.get('about', {}).get('campus_size', '20 acres'),
        'institute_type': 'Private Unaided (A+ Grade)',
        'academic_year': admission.get('academic_year', '2025-26'),
        'general_info': admission.get('general_info', {}),
        'contact': {
            'phone': location.get('phone', '02186-250146'),
            'email': location.get('email', 'principal@sknscoe.ac.in'),
            'website': location.get('website', 'https://www.sknscoe.ac.in'),
            'address': location.get('address', 'Korti, Pandharpur, Solapur'),
        },
        'branches': branches,
        'pg_courses': admission.get('pg_courses', []),
        'eligibility': admission.get('eligibility', {}),
        'admissionProcess': process_steps,
        'requiredDocuments': required_docs,
        'scholarshipSchemes': admission.get('scholarship_schemes', []),
        'placement': {
            'averagePackage': '5.2 LPA',
            'highestPackage': highest,
            'placedStudents': str(placement_data.get('total_students_placed_listed', 238)),
            'topRecruiters': placement_data.get('major_recruiters', []),
        },
        'note': admission.get('note', ''),
    }

    return jsonify({'success': True, 'data': data})


@app.route('/api/admission/fees', methods=['GET'])
def admission_fees():
    """Return fee structure data from Admission.json."""
    admission = _load_json_data('Admission.json')
    fee_structure = admission.get('fee_structure', {})
    cap = fee_structure.get('cap_seats_2025_26', {})

    # Build fee structure per branch (same fees apply to all UG branches under CAP)
    branches = {}
    for course in admission.get('ug_courses', []):
        code = course['course_name'].split()[0].upper()
        branches[code] = {
            'branch': course['course_name'],
            'cap_seats': cap.get('first_year_btech', []),
            'institute_level': fee_structure.get('institute_level_seats_2025_26', {}).get('first_year_btech', {}),
            'hostel': cap.get('hostel_and_mess_fees', {}),
        }

    data = {
        'feeStructure': branches,
        'hostelFees': cap.get('hostel_and_mess_fees', {}),
        'cautionDeposit': {
            'ug': cap.get('caution_deposit_ug', 2000),
            'pg': cap.get('caution_deposit_pg', 3000),
        },
        'note': admission.get('note', ''),
        'cap_fees_ug': cap.get('first_year_btech', []),
        'institute_level_ug': fee_structure.get('institute_level_seats_2025_26', {}).get('first_year_btech', {}),
    }
    return jsonify({'success': True, 'data': data})


@app.route('/api/admission/fees/calculate', methods=['GET'])
def admission_calculate_fees():
    """Calculate estimated fees based on branch, hostel, transport options."""
    branch = request.args.get('branch', 'CSE')
    year = int(request.args.get('year', 1))
    include_hostel = request.args.get('includeHostel', 'false').lower() == 'true'
    include_transport = request.args.get('includeTransport', 'false').lower() == 'true'
    category = request.args.get('category', 'OPEN (NON-EBC)')

    admission = _load_json_data('Admission.json')
    cap = admission.get('fee_structure', {}).get('cap_seats_2025_26', {})
    fee_rows = cap.get('first_year_btech', [])

    # Find matching category fee
    tuition = 83478.0
    development = 12522.0
    total_base = 96000.0

    for row in fee_rows:
        if row.get('category', '').upper() in category.upper():
            tuition = row.get('tuition_fees', tuition)
            development = row.get('development_fees', development)
            total_base = row.get('total_fees', total_base)
            break

    caution = cap.get('caution_deposit_ug', 2000.0) if year == 1 else 0
    hostel = cap.get('hostel_and_mess_fees', {}).get('total', 45708.0) if include_hostel else 0
    transport = 15000.0 if include_transport else 0

    breakdown = {
        'tuitionFees': tuition,
        'developmentFees': development,
        'cautionDeposit': caution,
        'hostelFees': hostel,
        'transportFees': transport,
    }
    total = sum(breakdown.values())

    return jsonify({
        'success': True,
        'calculation': {
            'branch': branch,
            'year': year,
            'category': category,
            'breakdown': breakdown,
            'totalFees': total
        }
    })


@app.route('/api/admission/contacts', methods=['GET'])
def admission_contacts():
    """Return contact information from Contact-Us.json and Institute Info.json."""
    contact_data = _load_json_data('Contact-Us.json')
    institute = _load_json_data('Institute Info.json')

    leadership = institute.get('leadership', {})
    dept_contacts = contact_data.get('department_contacts', [])

    # Admission office from contact data
    admission_office = {
        'phone': dept_contacts[3].get('phone_number', '+91 8275206048') if len(dept_contacts) > 3 else '+91 02186-250146',
        'email': contact_data.get('address', {}).get('email', 'principal@sknscoe.ac.in'),
        'address': contact_data.get('address', {}).get('full_address', 'SKN SKNSCOE, Korti, Pandharpur'),
        'workingHours': f"{contact_data.get('timings', {}).get('monday_to_friday', '9AM-5PM')} (Mon-Sat)",
    }

    # Branch coordinators mapped from department_contacts
    branch_coordinators = {}
    for entry in dept_contacts:
        dept = entry.get('department', '')
        phone = entry.get('phone_number', '')
        branch_coordinators[dept] = {'name': dept, 'phone': phone, 'email': 'contact@sknscoe.ac.in'}

    # College leadership
    principal = leadership.get('principal', {})

    return jsonify({
        'success': True,
        'data': {
            'admissionOffice': admission_office,
            'branchCoordinators': branch_coordinators,
            'principal': {
                'name': principal.get('name', 'Dr. Kailash J. Karande'),
                'title': principal.get('title', 'M.Tech, Ph.D'),
                'email': principal.get('contact', 'principal@sknscoe.ac.in'),
            },
            'website': institute.get('location', {}).get('website', 'https://www.sknscoe.ac.in'),
        }
    })


@app.route('/api/admission/all-data', methods=['GET'])
def admission_all_data():
    """Return all college JSON data as a single object for the chatbot/frontend."""
    all_data = {}
    for key, filename in [
        ('admission', 'Admission.json'),
        ('institute_info', 'Institute Info.json'),
        ('placements', 'Placements.json'),
        ('facilities', 'Facilities.json'),
        ('library', 'Library.json'),
        ('contact', 'Contact-Us.json'),
        ('research', 'Research.json'),
        ('about', 'About-Us.json'),
    ]:
        all_data[key] = _load_json_data(filename)
    return jsonify({'success': True, 'data': all_data})


# ============================================================
# PUBLIC CHATBOT ENDPOINT (No Auth Required — for Login page)
# ============================================================

@app.route('/api/chatbot/public', methods=['POST'])
def public_chatbot():
    """Handle chatbot queries from unauthenticated users (login page / admission queries)."""
    payload = request.get_json(silent=True) or {}
    message = (payload.get('message') or '').strip()

    if not message:
        return jsonify({'error': 'message is required'}), 400

    msg = message.lower()
    from nlp_engine import nlp_engine

    # Load college data for context
    admission = _load_json_data('Admission.json')
    contact = _load_json_data('Contact-Us.json')
    institute = _load_json_data('Institute Info.json')
    placements = _load_json_data('Placements.json')

    # --- Keyword-based direct answers (no DB, no auth needed) ---
    if any(w in msg for w in ['hello', 'hi', 'hey', 'namaste']):
        return jsonify({
            'response': "👋 Hello! I'm SmartEdu Assistant for SKN Sinhgad College of Engineering. I can help you with:\n\n• 🎓 Admission process & eligibility\n• 💳 Fee structure for 2025-26\n• 🏫 Available branches & courses\n• 📋 Required documents\n• 📞 Contact information\n• 🏆 Placements & scholarships\n\nWhat would you like to know?"
        })

    if any(w in msg for w in ['fee', 'fees', 'tuition', 'cost', 'charge', 'payment']):
        cap = admission.get('fee_structure', {}).get('cap_seats_2025_26', {})
        fee_rows = cap.get('first_year_btech', [])
        response = "💳 **Fee Structure 2025-26 (B.Tech First Year — CAP Seats):**\n\n"
        for row in fee_rows:
            response += f"• **{row.get('category')}**: ₹{row.get('total_fees', 0):,.0f}/year\n"
        hostel = cap.get('hostel_and_mess_fees', {})
        response += f"\n🏠 **Hostel + Mess Fees**: ₹{hostel.get('total', 45708):,.0f}/year\n"
        response += "\n*Institute Level Seats: ₹2,88,000/year*\n\nFor OBC/SC/ST/NT categories, fees are substantially reduced or waived. Ask me about scholarships!"
        return jsonify({'response': response})

    if any(w in msg for w in ['branch', 'course', 'department', 'available', 'programme', 'program']):
        ug = admission.get('ug_courses', [])
        pg = admission.get('pg_courses', [])
        response = "🎓 **Available Courses at SKNSCOE:**\n\n**Undergraduate (B.Tech):**\n"
        for c in ug:
            response += f"• {c['course_name']} — {c['intake']} seats\n"
        response += "\n**Postgraduate (M.Tech):**\n"
        for c in pg[:3]:
            response += f"• {c['course_name']} — {c['intake']} seats\n"
        response += "\nAll programs are AICTE approved and affiliated to Solapur University."
        return jsonify({'response': response})

    if any(w in msg for w in ['eligib', 'qualify', 'require', 'criteria', 'marks', 'percentage', 'cutoff']):
        response = "📋 **Eligibility for B.Tech First Year (2025-26):**\n\n"
        for criterion in admission.get('eligibility', {}).get('first_year_ug', {}).get('maharashtra_state', [])[:4]:
            response += f"✅ {criterion}\n\n"
        response += "**Entrance Exam:** Valid MHT-CET or JEE Main score required.\n\n**OBC/Reserved categories:** Minimum 40% marks in PCM."
        return jsonify({'response': response})

    if any(w in msg for w in ['document', 'certificate', 'required', 'submit', 'bring', 'original']):
        docs = ['HSC (12th) Marksheet', 'SSC (10th) Marksheet', 'MHT-CET Score Card', 'Transfer Certificate', 'Caste Certificate (if applicable)', 'Aadhaar Card', '4 Passport Photos', 'DTE Allotment Letter']
        response = "📄 **Required Documents for Admission:**\n\n" + "\n".join(f"✅ {d}" for d in docs)
        response += "\n\n*Please bring originals + 2 attested photocopies of each.*"
        return jsonify({'response': response})

    if any(w in msg for w in ['scholarship', 'tfws', 'ebc', 'obc', 'sc', 'st', 'free', 'waiver', 'financial']):
        response = "🎁 **Scholarships & Fee Waivers:**\n\n"
        for scheme in admission.get('scholarship_schemes', []):
            response += f"**{scheme.get('department')}:**\n"
            for s in scheme.get('schemes', []):
                response += f"  • {s}\n"
            response += "\n"
        response += "TFWS (Tuition Fee Waiver Scheme): 100% tuition fee waiver for eligible students!\n\nContact Accounts Section: +91 9423536913"
        return jsonify({'response': response})

    if any(w in msg for w in ['placement', 'job', 'salary', 'package', 'company', 'recruiter', 'career']):
        stats = placements.get('placement_statistics_2023_24', {})
        recruiters = [r['company'] for r in placements.get('top_recruiters', [])[:8]]
        response = f"🏆 **Placement Statistics (2023-24):**\n\n"
        response += f"• Average Package: {stats.get('average_package', '5.2 LPA')}\n"
        response += f"• Highest Package: {stats.get('highest_package', '18 LPA')}\n"
        response += f"• Students Placed: {stats.get('placed_students', '85%')}\n\n"
        response += f"🏢 **Top Recruiters:** {', '.join(recruiters)}\n\n"
        response += "TPO Contact: +91 8308614875"
        return jsonify({'response': response})

    if any(w in msg for w in ['contact', 'phone', 'address', 'office', 'reach', 'location']):
        depts = contact.get('department_contacts', [])[:5]
        timings = contact.get('timings', {})
        response = "📞 **Contact SKNSCOE:**\n\n"
        response += f"📍 **Address:** {contact.get('address', {}).get('full_address', 'Korti, Pandharpur')}\n"
        response += f"✉️ **Email:** {contact.get('address', {}).get('email', 'principal@sknscoe.ac.in')}\n\n"
        response += "**Key Contacts:**\n"
        for d in depts:
            response += f"• {d['department']}: {d['phone_number']}\n"
        response += f"\n⏰ **Working Hours:** {timings.get('monday_to_friday', '9AM-5PM')} (Mon-Fri)"
        return jsonify({'response': response})

    if any(w in msg for w in ['hostel', 'accommodation', 'room', 'boarding', 'stay', 'mess']):
        cap = admission.get('fee_structure', {}).get('cap_seats_2025_26', {})
        hostel = cap.get('hostel_and_mess_fees', {})
        response = f"🏠 **Hostel & Accommodation:**\n\nSKNSCOE is a **fully residential campus** with hostel facilities.\n\n"
        response += f"• Hostel + Mess Fee: ₹{hostel.get('hostel_mess_fee', 42000):,.0f}/year\n"
        response += f"• Caution Deposit: ₹{hostel.get('caution_money_deposit', 3000):,.0f} (refundable)\n"
        response += f"• Admission Prospectus: ₹{hostel.get('hostel_admission_form_prospectus', 708):,.0f}\n"
        response += f"• **Total: ₹{hostel.get('total', 45708):,.0f}/year**\n\n"
        response += "The campus includes mess, canteen, dispensary, ATM, gymnasium, and sports facilities."
        return jsonify({'response': response})

    if any(w in msg for w in ['process', 'apply', 'how to', 'steps', 'procedure', 'admission']):
        response = "📋 **Admission Process 2025-26:**\n\n"
        steps = [
            "1️⃣ Appear in **MHT-CET 2025** (or JEE Main) — mandatory for UG admissions",
            "2️⃣ Register on **DTE Maharashtra portal** (dtemaharashtra.gov.in)",
            "3️⃣ Fill **Choice Form** — select SKNSCOE branches as preference",
            "4️⃣ Check **CAP Allotment** result",
            "5️⃣ **Confirm seat** by paying acceptance fee online",
            "6️⃣ **Report to college** with original documents for verification",
            "7️⃣ Pay **annual fees** and collect enrollment confirmation",
        ]
        response += "\n".join(steps)
        response += "\n\n💡 *Tip: For Lateral Entry (DSE) — Diploma holders can apply directly to 2nd year!*"
        return jsonify({'response': response})

    # Fallback: use NLP engine with public context
    user_context = {'role': 'guest', 'name': 'Prospective Student'}
    student_db_context = f"This is a prospective student asking about admission to SKNSCOE."

    try:
        response = nlp_engine.generate_gemini_response(
            message,
            intent='ask_admission',
            user_context=user_context,
            student_db_context=student_db_context,
        )
        return jsonify({'response': response})
    except Exception as e:
        return jsonify({'response': "I can help you with admission info! Try asking about:\n• Fee structure\n• Available branches\n• Eligibility criteria\n• Required documents\n• Scholarship schemes\n• Hostel facilities\n• Placement statistics"})

# ==========================================
# AI Services Routes (Notes, Quiz, Guidance)
# ==========================================

def extract_text_from_file(file):
    """Extract text from PDF, DOCX, or TXT files."""
    filename = file.filename.lower()
    extracted = ''
    try:
        if filename.endswith('.pdf'):
            import io
            from pypdf import PdfReader
            file_bytes = file.read()
            if not file_bytes:
                return ''
            reader = PdfReader(io.BytesIO(file_bytes))
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    extracted += text + '\n'
        elif filename.endswith('.docx'):
            import io
            import docx
            file_bytes = file.read()
            if not file_bytes:
                return ''
            doc = docx.Document(io.BytesIO(file_bytes))
            for para in doc.paragraphs:
                extracted += para.text + '\n'
        elif filename.endswith('.txt'):
            extracted = file.read().decode('utf-8', errors='ignore')
        else:
            extracted = ''
    except Exception as e:
        print(f'CRITICAL: File extraction error ({filename}): {str(e)}')
        extracted = ''
    print(f'DEBUG: Extracted {len(extracted)} characters from {filename}')
    return extracted.strip()


@app.route('/api/ai/notes', methods=['POST'])
def generate_ai_notes():
    from nlp_engine import nlp_engine
    if not getattr(nlp_engine, 'groq_client', None):
        return jsonify({'error': 'AI engine not configured. Please check GROQ_API_KEY.'}), 503

    content = request.form.get('content', '')
    if 'file' in request.files:
        file = request.files['file']
        if file and file.filename:
            extracted = extract_text_from_file(file)
            if extracted:
                content = extracted + ('\n' + content if content else '')
            elif not content:
                return jsonify({'error': f'Could not extract text from {file.filename}. Ensure PDF has selectable text.'}), 400

    if not content.strip():
        return jsonify({'error': 'Please provide content as text or upload a PDF/DOCX/TXT file.'}), 400

    prompt = (
        "You are a smart study assistant. Generate comprehensive, well-structured study notes for the following content.\n"
        "Format rules:\n"
        "- Use ## for main headings, ### for sub-headings\n"
        "- Use bullet points (- ) for key concepts\n"
        "- Bold (**text**) important terms\n"
        "- Add a '## Key Takeaways' section at the end\n"
        "- Keep it concise yet complete\n\n"
        f"Content:\n{content[:6000]}"
    )
    try:
        response = nlp_engine.groq_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.4,
            max_tokens=1500
        )
        return jsonify({'notes': response.choices[0].message.content})
    except Exception as e:
        return jsonify({'error': f'AI generation failed: {str(e)}'}), 500


@app.route('/api/ai/quiz', methods=['POST'])
def generate_ai_quiz():
    from nlp_engine import nlp_engine
    if not getattr(nlp_engine, 'groq_client', None):
        return jsonify({'error': 'AI engine not configured. Please check GROQ_API_KEY.'}), 503

    content = request.form.get('content', '')
    num_questions = 5
    try:
        num_questions = max(1, min(20, int(request.form.get('num_questions', 5))))
    except:
        num_questions = 5

    if 'file' in request.files:
        file = request.files['file']
        if file and file.filename:
            extracted = extract_text_from_file(file)
            if extracted:
                content = extracted + ('\n' + content if content else '')

    if not content.strip():
        return jsonify({'error': 'Please provide content as text or upload a PDF/DOCX/TXT file.'}), 400

    prompt = (
        f"Create exactly {num_questions} high-quality, academic multiple-choice questions for engineering students based on the provided content.\n"
        "CRITICAL: Return ONLY a valid JSON array. No conversational text, no markdown blocks.\n"
        "Each question must have 4 unique, plausible options derived directly from the content.\n"
        "DO NOT use generic labels like 'Option A' or 'Correct Answer' as placeholders. Use actual subject-specific text.\n\n"
        "Example of desired output format:\n"
        "[\n"
        "  {\n"
        '    "question": "What is the primary purpose of significant indentation in Python?",\n'
        '    "options": ["To define code blocks", "To increase file size", "To slow down execution", "To allow for comments"],\n'
        '    "correct_answer": "To define code blocks"\n'
        "  }\n"
        "]\n\n"
        f"Generate exactly {num_questions} items using the same structure for the following content:\n"
        f"\n{content[:6000]}"
    )
    try:
        response = nlp_engine.groq_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.2,
            max_tokens=1500
        )
        reply = response.choices[0].message.content.strip()

        # Strip markdown code blocks if present
        if '```json' in reply:
            reply = reply.split('```json')[1].split('```')[0].strip()
        elif '```' in reply:
            reply = reply.split('```')[1].split('```')[0].strip()

        # Find JSON array boundaries
        start = reply.find('[')
        end = reply.rfind(']')
        if start != -1 and end != -1:
            reply = reply[start:end+1]

        questions = json.loads(reply)
        return jsonify({'questions': questions})
    except json.JSONDecodeError as e:
        return jsonify({'error': f'AI returned invalid JSON. Please try again. ({str(e)})'}), 500
    except Exception as e:
        return jsonify({'error': f'AI generation failed: {str(e)}'}), 500


@app.route('/api/ai/guidance', methods=['POST'])
def generate_ai_guidance():
    payload = request.get_json(silent=True) or {}
    user_id = payload.get('user_id')
    subject = payload.get('subject', '')
    response_text = get_personalized_guidance(user_id, subject)
    return jsonify({'guidance': response_text})


if __name__ == '__main__':


    # For local development only: create tables if they don't exist and run the app.
    # In production, use proper migrations and a WSGI server.
    with app.app_context():
        try:
            db.create_all()
            print('Database tables ensured (db.create_all)')
        except Exception as e:
            print(f'Warning: could not create tables automatically: {e}')

    # Start Flask development server
    app.run(host='127.0.0.1', port=5000, debug=app.config.get('DEBUG', False))


